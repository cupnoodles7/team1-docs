# -*- coding: utf-8 -*-
"""Build TF_Reader_Team1_Weekly_Plan.docx — deliverables, assumptions, dependencies. No names."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'final_plan_v1_backup.docx'   # inherit styles / page setup
OUT = 'TF_Reader_Team1_Weekly_Plan.docx'

NAVY, DEEP, INK, GOLD, TINT = '13487F', '0D3562', '16233B', 'C8922F', 'E9F1FB'
KICKER, SUBTLE, LABEL, WHITE = 'A9C9EE', 'D4E3F6', '8FB4DE', 'FFFFFF'

doc = Document(SRC)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for ch in list(body):
    if ch is not sectPr:
        body.remove(ch)

def para(sa=None, sb=None, style=None):
    p = doc.add_paragraph(style=style)
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    if sb is not None: p.paragraph_format.space_before = Pt(sb)
    return p

def run(p, t, size=None, bold=None, color=None, font=None):
    r = p.add_run(t)
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor.from_string(color)
    if font: r.font.name = font
    return r

def rich(p, parts):
    for x in parts:
        if isinstance(x, str): p.add_run(x)
        else: run(p, x[0], bold=x[1], color=INK if x[1] else None)

def secnum(n):
    run(para(sa=0), n, size=9, bold=True, color=GOLD, font='Consolas')

def sectitle(t):
    run(para(sa=8), t, size=19, bold=True, color=DEEP, font='Georgia')

def h3(t):
    run(para(sb=14, sa=4), t, size=13, bold=True, color=NAVY, font='Georgia')

def label(t):
    run(para(sb=11, sa=3), t, size=10.5, bold=True, color=INK)

def bp(parts, sa=6):
    p = para(sa=sa); rich(p, parts if isinstance(parts, list) else [parts]); return p

def bullet(parts):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p

def spacer(pts=8): para(sa=pts)

def shade(c, fill):
    tcPr = c._tc.get_or_add_tcPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto'); e.set(qn('w:fill'), fill)
    tcPr.append(e)

def ctr(t):
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); t._tbl.tblPr.append(jc)
    for r in t.rows:
        trPr = r._tr.get_or_add_trPr()
        j = OxmlElement('w:jc'); j.set(qn('w:val'), 'center'); trPr.append(j)

def w(c, tw):
    tcPr = c._tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcW')): tcPr.remove(o)
    e = OxmlElement('w:tcW'); e.set(qn('w:w'), str(tw)); e.set(qn('w:type'), 'dxa'); tcPr.insert(0, e)

def ctext(c, txt, size=9, bold=False, color=INK):
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run(p, txt, size=size, bold=bold, color=color)

def table(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; ctr(t)
    for i, htx in enumerate(headers):
        c = t.rows[0].cells[i]; w(c, widths[i]); shade(c, NAVY)
        ctext(c, htx, size=8.5, bold=True, color=WHITE)
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            c = t.rows[ri].cells[ci]; w(c, widths[ci])
            ctext(c, v, size=9, bold=(ci == 0), color=NAVY if ci == 0 else INK)
    spacer(8); return t

def callout(kick, text):
    t = doc.add_table(rows=1, cols=1); t.style = None; ctr(t)
    c = t.rows[0].cells[0]; w(c, 9792); shade(c, TINT)
    tcPr = c._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for side, sz, col in (('top',4,TINT),('left',24,NAVY),('bottom',4,TINT),('right',4,TINT)):
        e = OxmlElement('w:'+side)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz)); e.set(qn('w:space'),'0'); e.set(qn('w:color'),col)
        b.append(e)
    tcPr.append(b)
    p1 = c.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
    run(p1, kick, size=8.5, bold=True, color=NAVY)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    run(p2, text, size=9.5)
    spacer(8)

# ------------------------------------------------------------------ HERO
hero = doc.add_table(rows=1, cols=1); hero.style = None
c = hero.rows[0].cells[0]; w(c, 9792); shade(c, NAVY)
p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(4)
run(p, 'WEEKLY DELIVERY PLAN  ·  DISCOVERY & SELECTION', size=9, bold=True, color=KICKER)
p = c.add_paragraph(); p.paragraph_format.space_after = Pt(6)
run(p, 'What we deliver, week by week', size=27, bold=True, color=WHITE, font='Georgia')
p = c.add_paragraph(); p.paragraph_format.space_after = Pt(18)
run(p, 'Deliverables per week, the assumptions they rest on, and what we need from other teams to hold the dates.',
    size=11.5, color=SUBTLE)
for k, v in (('PROGRAMME   ', 'T&F Reader — 8-Week Graduate Prototype'),
             ('WE OWN   ', 'Institution listing · Institution selection · Catalogue browse · Unified search'),
             ('TARGET   ', 'Feature-complete on mock data by end of Week 3 · integration closes Week 4'),
             ('WEEKS   ', 'W1 10–14 Aug · W2 17–21 Aug · W3 24–28 Aug · W4 31 Aug–4 Sep 2026')):
    p = c.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, k, size=8, bold=True, color=LABEL); run(p, v, size=9.5, color=WHITE)
p = c.add_paragraph(); p.paragraph_format.space_after = Pt(16)
spacer(10)

# =========================================================== 01 WHAT WE OWN
secnum('01'); sectitle('What we are building')
bp('Six areas. Everything in them is working against mock data by the end of Week 3; integration with the '
   'real services runs through Week 4.')
table(['Area', 'What it covers'],
 [['Catalogue', 'Browse the institution\'s catalogue as shelves, read from one root feed and shown as tabs, each paginated by following its next link. Open access only before sign-in; the institution\'s entitled catalogue after.'],
  ['Search & filter', 'A search surface over the catalogue team\'s server-side search. Filters for file format, access tier, subject, publisher and year. Sort, result count, and voice input. The matching and ranking are theirs, not ours.'],
  ['Institutions', 'Browse and search institutions over their paged endpoint, view detail, select one, remember it, change or clear it, and hand off to sign-in.'],
  ['Access & actions', 'One rule deciding the single action on every item — Read, Download, Borrow, Sign in or Waitlist — read from the acquisition link the catalogue team send, plus the loan that Borrow creates.'],
  ['Item detail', 'Full metadata for books, journals, articles and audio, with the same action model as the list, plus the live copy count that only this screen can fetch.'],
  ['App & states', 'Shell, navigation, shared components, and the loading, empty, error and offline states across every surface.']],
 [1500, 8292])

callout('WHAT WE ARE NOT BUILDING',
        'The reader itself and the downloads library belong to the reader team. Authentication screens and '
        'the borrow service belong to the authentication team. The catalogue feeds, the institution data '
        'AND THE SEARCH ITSELF belong to the catalogue team. We call all three; we do not build them. '
        'Purchase, Subscribe and individual subscriber accounts are all out of scope for the prototype '
        'entirely — the catalogue team cut individual subscribers at their own design gate, because their '
        'entitlement model is keyed by institution.')

# ======================================================== 02 ASSUMPTIONS
secnum('02'); sectitle('What we are assuming')
bp([('MOST OF THESE ARE NOW CONFIRMED OR CORRECTED.', True),
    ' The catalogue team have published a full technical specification, and it answers or removes the '
    'premise of most of the list below. Two of the assumptions were WRONG in a way that mattered, and the '
    'design document has been corrected as a result. The status column says which is which. Everything '
    'still open is built against a stub, so nothing is blocked.'])

table(['#', 'Assumption', 'Status now'],
 [['1', 'The catalogue team can serve a full catalogue to a user who is not signed in', 'WRONG, AND NARROWER. There is an unauthenticated feed, but it carries OPEN ACCESS TITLES ONLY. The full catalogue is entitlement-scoped and token-protected. The browse-first model survives, but what an anonymous user sees is much smaller than assumed. This reverses a statement in the signed design document, and it is now a fact rather than a proposal.'],
  ['2', 'They will add a per-item field telling us the access tier', 'THERE IS NO FIELD, AND THERE WILL NOT BE ONE — but we do not need it. Confirmed 11 August: their earlier document showed such a field and that document is out of date. What they do send is the licence model on the borrow link, and they have given us the mapping — unlimited means subscription, concurrent means premium, and no licence model at all means open access. We work the tier out once when the data arrives, so every screen still receives it as a plain value. The longest-lead request in the plan closed in a single message, just not the way anyone expected.'],
  ['3', 'After sign-in we receive a feed already scoped to entitlements', 'CONFIRMED. They scope it. We do not duplicate their entitlement logic.'],
  ['4', 'The three feeds are split by content type and shown as three tabs', 'THE PREMISE WAS WRONG. There are not three feeds. There is one root feed per institution carrying shelves, and the shelf set is CONFIGURED PER INSTITUTION by their admin console. Their content-type field is a FILE FORMAT — PDF, EPUB, audio — and it does not divide the catalogue. Shelves still show as tabs, and the tab bar must be driven from data because they can re-cut shelves without a rebuild.'],
  ['5', 'Articles sit inside the Journals feed, and that feed is flat', 'MOOT — there is no Journals feed. See assumption 4, and the new gap at 14.'],
  ['6', 'Search matching, ranking, sorting and filtering are ours to build', 'WRONG, AND IT IS A REDUCTION. Catalogue search is SERVER-SIDE, because results are filtered by what the institution is entitled to and a member must never see a result they cannot open. Matching, ranking and tokenisation are all theirs. We build the query surface. This is the largest scope reduction in the plan — but their endpoint arrives in Week 4, which is our integration week, so it is a schedule risk in exchange.'],
  ['7', 'Institution data has no OPDS equivalent, so we define its shape', 'THEY DEFINED IT, and published it in full. We adopt theirs. Their endpoint also arrives in Week 2, earlier than anything else we depend on.'],
  ['8', 'A loan is needed for anything the user can take a copy of. Open Access and Subscription need Borrow first; Elite is read-only so never does', 'WRONG ON TWO OF THREE TIERS, AND THIS IS THE MOST CONSEQUENTIAL CORRECTION IN THIS DOCUMENT. The real rule is that a loan exists for anything ENCRYPTED — Subscription and Elite — and a queue exists only where copies are FINITE, which is Elite alone. Open access needs no loan in ANY state, including signed in. Elite is the ONE tier that borrows, consumes a copy and can queue — and the one tier that cannot be downloaded. The signed design document has been corrected and needs ratifying.'],
  ['9', 'The authentication team can tell us, in one call, which of a page of items a user has on loan', 'STILL OPEN, and now the highest-risk item on the list. Forty results would otherwise mean forty calls. The catalogue team ship exactly this shape for item metadata, capped at 100 ids, which is the precedent to cite.'],
  ['10', 'A signed-out user reading Open Access needs no loan', 'CONFIRMED, and it goes further: a SIGNED-IN user reading open access needs no loan either. Open access is stored unencrypted precisely so anyone can open it.'],
  ['11', 'Subscribe applies only to individual accounts, never institutional ones', 'REMOVED ENTIRELY. Individual subscribers are cut from scope at the catalogue team\'s design gate. Subscribe is deleted from the action vocabulary, and the personal-account option on our sign-in choice screen becomes a browse-free-content entry instead.'],
  ['12', 'We hold the item a user tapped across the whole sign-in journey and return them to it', 'STILL OURS, and still true. Narrower in one way: sign-in is ALWAYS SAML, so there is no auth-type branch to build.'],
  ['13', 'Sample data arrives in Week 1 and real endpoints in Week 3', 'COMMITTED for Week 1 — an API specification, three sample feeds and mock endpoints. This is the single most valuable thing we receive all project, because it closes our largest internal risk. Real endpoints arrive across Weeks 2 to 4 rather than all in Week 3.'],
  ['14', 'NEW GAP — the catalogue team model no WORK TYPE. Their content-type field is a file format. But our article-detail and book-detail screens differ by exactly that axis', 'MOSTLY CLOSED, 11 August. Their sample feeds label every title with a standard schema.org type, and they have confirmed it is official — so book and audiobook are settled. Only the journal and article labels are still to come, and they affect which detail screen renders, which is Week 2 work. We carry two marked guesses meanwhile and fall back safely on anything we do not recognise.'],
  ['15', 'NEW GAP — there is no DOI field anywhere in their schema', 'CLOSED, 11 August — there is no DOI and there will not be one. Only their own identifier exists. DOI comes off the article detail screen, and that is now a design change to make rather than an answer to wait for. Also corrected: ISBN is NOT the plain field their document showed. In the actual feed it is wrapped inside a longer identifier string that has to be unpacked, and titles without an ISBN carry an internal identifier instead.'],
  ['16', 'NEW GAP — their document lists us as owning the downloads library screen; our plan assigns it to the reader team', 'A whole screen and one component. Being resolved in Week 1; escalated if it is not.'],
  ['17', 'NEW — subject filters exist ONLY IF WE ASK FOR THEM. Subjects are dynamic, so they cannot be known per institution without extra work on their side', 'Their stated default if we say nothing is not to build it — which would delete the Browse-by-Subject row from the home screen. We are answering yes in Week 1.']],
 [400, 4400, 4992])

# ======================================================= 03 WEEK BY WEEK
secnum('03'); sectitle('Week by week')
callout('HOW TO READ THIS',
        'Weeks 1 to 3 are ours alone and run entirely on mock data, so nothing in them depends on another '
        'team arriving on time. Week 4 is integration, and is the only week whose content depends on '
        'services outside the team. It is deliberately left light.')

h3('Week 1 — 10 to 14 August · Foundations')
label('Delivers')
bullet('Design tokens, shared component library, app shell and navigation — all merged before feature work starts')
bullet('The internal content model, and adapters turning all three OPDS feeds into it')
bullet('Identifier handling, so DOI and ISBN become searchable fields')
bullet('Mock data covering every content type and every access state')
bullet('Catalogue browse with feed tabs and pagination')
bullet('Search with matching and ranking, working over mock data')
bullet('Institution list, detail and search')
label('Needs from others')
bullet('Sample institution list and detail — the catalogue samples have arrived, these have not, and they are the whole of one person\'s week')
bullet('Confirmation that an unsigned-in user can fetch the full catalogue')
bullet('Whether access tier can still be used as a filter, now that there is no tier field to filter on')
bullet('The sign-in handoff contract, even if the screen lands later')
callout('DONE WHEN', 'The catalogue is browsable and searchable on mock data, institutions can be searched '
        'and viewed, and every contract request has gone out in writing.')

h3('Week 2 — 17 to 21 August · Access, loans and detail')
label('Delivers')
bullet('Institution selection, remembered across restarts, changeable and clearable')
bullet('Handing off to sign-in with the institution\'s identity-provider hint. Always SAML, so there is no route to choose')
bullet('The access rule producing exactly one action per item, read from the acquisition link the catalogue team send. The tier badge on every card is what the tier is now FOR')
bullet('Borrow creating a loan on Subscription and Elite. The button then shows Read and Download on Subscription, and Read ALONE on Elite')
bullet('Loan state fetched and cached, so the right action shows on a full page of results')
bullet('The tapped item held across sign-in and restored afterwards')
bullet('Item detail for books, journals, articles and audio, plus the live copy count on Elite titles, which only this screen can fetch')
bullet('Filters, sort, result count, and search state that survives navigation')
bullet('Voice input feeding the same search')
label('Needs from others')
bullet('A way to read the signed-in user’s session and entitlements')
bullet('A single call returning loan state for a page of items')
bullet('The Borrow endpoint, or its contract so we can build against it')
bullet('A sign-in screen we can route into, even a mock one')
callout('DONE WHEN', 'BOTH states — anonymous and institutional — can be walked end to end on mock data, '
        'with the correct single action on every item and the return-to-item journey working. The individual '
        'state is deleted. Watch specifically for Elite showing ONE button and Subscription showing TWO: '
        'that is the correction most likely to be got wrong.')

h3('Week 3 — 24 to 28 August · Hardening, and the first real data')
label('Delivers')
bullet('Loading, empty, error and offline states across every surface')
bullet('Waitlist and access-restricted states, and Borrow failure handling driven by the catalogue team\'s enumerated error codes, so a denial can say the library subscription has expired rather than not available')
bullet('Every screen checked against every state — the largest single block of the week')
bullet('Search hardening. The access-tier filter has moved forward to Week 2, because the field already exists')
bullet('Institution list working offline, since sign-in depends on it')
bullet('Real feeds and real sessions wired in as they become available')
label('Needs from others')
bullet('Real catalogue and institution endpoints')
bullet('Real session and entitlement access')
callout('DONE WHEN — THE ONE THAT MATTERS', 'Everything we own is finished, hardened and demonstrable. '
        'From this point anything outstanding is another team’s service, not our feature work.')

h3('Week 4 — 31 August to 4 September · Integration')
label('Delivers')
bullet('Mock data swapped for real catalogue, session, loan and reader services')
bullet('Fixes for any differences between the assumed and real data shapes')
bullet('Accessibility pass and the behaviour-driven test suite')
bullet('Buffer for anything that arrived late')
callout('IF NOTHING LANDS', 'This week is deliberately underfilled. If everything arrives on time we finish '
        'ahead. If nothing does, we demonstrate on mock data and say so plainly, rather than discovering '
        'the gap on the day.')

# ====================================================== 04 DEPENDENCIES
secnum('04'); sectitle('What we need, and by when')
bp('Grouped by who we need it from. Everything here already has a stub behind it, so a late answer costs '
   'us reshaping rather than idle time — except the four marked critical, which have no workable substitute.')

h3('What the catalogue team need FROM US — and three are due in Week 1')
bp([('This table did not exist before, and it is the one with our name and a date on it.', True),
    ' Each has a stated default if we stay silent, and on subject filters that default removes a feature '
    'the design already specifies. Silence is not neutral here.'])
table(['They need', 'By', 'What happens if we say nothing'],
 [['Which fields the institution list screen needs, and how we want it sorted', 'Wk 1', 'They ship a superset and refine. Our answer: their current shape is already right; sort alphabetically by name'],
  ['Whether we want subject filters', 'Wk 1', 'THEY DO NOT BUILD THEM — which deletes the Browse-by-Subject row from the home screen. Our answer is yes'],
  ['Confirmation that we render the acquisition link rather than deciding buttons ourselves', 'Wk 1', 'They write it as a test on their side. Our answer: confirmed'],
  ['Whether the home screen offers a browse-free-content entry beside find-your-institution', 'Wk 2', 'They ship the public feed anyway and it costs us a button. Our answer: yes, and it replaces the personal-account option']],
 [5100, 600, 4092])

h3('From the catalogue team')
table(['Need', 'By', 'Why it matters'],
 [['DELIVERED — three sample feeds, 10 August', 'done', 'The highest-value item on any of these lists, and it has landed. It closed our largest internal risk and corrected seven things in our own contract file. Their written specification is now OUT OF DATE by their own statement: where it disagrees with a sample, the sample wins'],
  ['Sample shapes we still do not have — the empty-search result, an error response, and the institution list and detail', 'Wk 1', 'Promised. The institution one is the whole of Keshav\'s block and those endpoints land in Week 2, so chase that one first'],
  ['The schema.org labels for journal and article', 'Wk 2', 'Book and audiobook are confirmed. These two decide which detail screen renders, which is Week 2 work. Two marked guesses in the meantime'],
  ['Whether access tier still works as a filter, now that there is no tier field', 'Wk 1', 'The filter chips were designed on the basis that it was free. Moktik builds them on Day 3, so this is a same-week answer'],
  ['Subject facets, so the subject filter and Browse-by-Subject can exist', 'Wk 1', 'They build it only if we ask. Roughly half a day of theirs'],
  ['The real institution directory', 'Wk 2', 'The earliest real dependency we receive from anyone'],
  ['The real entitlement-scoped catalogue feed and publication detail', 'Wk 3', 'We hold on fixtures until it lands'],
  ['Catalogue search, shelf paging, the public feed and the batch metadata endpoint', 'Wk 4', 'Arrives in our integration week, which is a schedule risk we are carrying deliberately'],
  ['ANSWERED — the access tier (derived from the licence model, no field), work type, DOI (there is none), publication date, one format per title, pagination, institution schema, logos, sign-in type, the error model, feed scoping', 'done', 'Do not ask again. Note the first three closed differently from how the plan assumed they would']],
 [5100, 600, 4092])

h3('From the authentication team')
table(['Need', 'By', 'Why it matters'],
 [['One call returning loan state for a page of items', 'Wk 1', 'CRITICAL — per-item calls would make a page of results visibly slow'],
  ['The sign-in handoff contract — what we pass, how control returns', 'Wk 1', 'We route to a placeholder otherwise'],
  ['Whether sign-in leaves the app, and how control comes back', 'Wk 1', 'Decides whether we store the tapped item on disk'],
  ['REMOVED — whether the session distinguishes institutional from individual. There are no individual accounts', '—', 'Individual subscribers are cut from scope'],
  ['Confirmation of the CORRECTED loan model: open access needs no loan ever; Subscription writes a loan and skips the copy counter; Elite takes a copy or writes a queue position', 'Wk 1', 'We had this backwards until the catalogue team published their specification. It changes the button on most of the catalogue'],
  ['How we read the signed-in user’s session. We do NOT need entitlement data — the catalogue team scope the feed themselves', 'Wk 2', 'We mock the shape otherwise'],
  ['The Borrow endpoint and its full set of failure responses, with the enumerated reason codes passed through unchanged', 'Wk 2', 'The codes are what let our error messages be specific rather than generic'],
  ['The live copy count for Elite titles, so we can show a waitlist before the user taps', 'Wk 4', 'The feed carries a total but never a live figure, so without this the detail screen shows no count'],
  ['Loan duration, and what happens when one lapses mid-read', 'Wk 2', 'We treat a lapsed loan as no loan'],
  ['Session expiry behaviour while browsing', 'Wk 2', 'Handled generically otherwise'],
  ['A mock sign-in screen we can route into', 'Wk 2', 'Placeholder otherwise']],
 [5100, 600, 4092])

h3('From the reader team')
table(['Need', 'By', 'Why it matters'],
 [['How we launch the reader from an item', 'Wk 2', 'Stubbed otherwise'],
  ['How a download is handed over, and who stores it', 'Wk 2', 'Stubbed otherwise'],
  ['Confirmation that Elite content cannot be downloaded at your end either', 'Wk 2', 'We simply never offer it'],
  ['WHO OWNS THE DOWNLOADS LIBRARY SCREEN. The catalogue team\'s document lists it as ours; our plan assigns it to the reader team with the reader itself', 'Wk 1', 'A whole screen and one component. If it is ours, one person\'s week is re-cut. Escalated if unresolved by Friday'],
  ['Who shows the downloaded indicator on catalogue items', 'Wk 2', 'Omitted initially'],
  ['Whether signed-out users can download, and who stores it', 'Wk 2', 'Omitted otherwise. Note this is the one case where it is possible at all — open access is unencrypted']],
 [5100, 600, 4092])

h3('From leadership')
table(['Need', 'By', 'Why it matters'],
 [['RATIFICATION OF THE CORRECTED ACCESS RULES in the signed design document. Its table is inverted on two of the three access tiers: it says open access needs a loan once signed in, and that the premium tier never borrows. The catalogue team\'s specification says the opposite on both, and we believe theirs', 'Wk 1', 'CRITICAL, AND THE MOST URGENT ITEM ON THIS PAGE. We have corrected the document and are building to the correction. Until it is ratified there are two contradictory rule sets in circulation and nobody can build against a single agreed source'],
  ['Sign-off that the catalogue an unsigned-in user sees is OPEN ACCESS ONLY, which reverses the signed design document', 'Wk 1', 'This is no longer a change we are proposing — the catalogue team cannot serve the original model. Send it with the item above; they are one conversation'],
  ['Confirmation that catalogue browse and search are ours to build', 'Wk 1', 'CRITICAL — most of this plan assumes it. The catalogue team\'s own document independently lists it as ours, which supports it'],
  ['Confirmation that purchase, Subscribe and individual subscriber accounts are all out of scope', 'Wk 1', 'The catalogue team have already cut individual subscribers at their own gate. This removes an action from every screen and changes our sign-in choice screen'],
  ['Sign-off that the catalogue presents as shelf tabs rather than one list', 'Wk 1', 'Still differs from the design mockups, though narrower than before: the shelf set is the institution\'s to configure, so there is nothing to merge or split'],
  ['An owner for the downloads library screen, if the two teams cannot settle it', 'Wk 1', 'A whole screen. Only comes to leadership if unresolved by Friday'],
  ['One shared application across teams, or one each', 'Wk 1', 'We build our own and keep it portable. The catalogue team confirm we own no backend at all']],
 [5100, 600, 4092])

callout('THE FOUR THAT DECIDE THE DATES — REVISED',
        'Three of the previous five are answered, and the list is now different. The Week 1 sample data and '
        'API specification, because it closes our largest internal risk and a Monday phone call is what '
        'gets it. The work-type field, because two of our screens differ by an axis the backend does not '
        'model. Ratification of the corrected access rules, because two contradictory versions are '
        'currently in circulation. And a single call for loan state across a page of results, which is the '
        'one remaining item with no workable substitute. The access tier turned out to need no field at all — '
        'we work it out from the licence model they already send, using their own mapping. '
        'The unauthenticated feed exists but is smaller than assumed. Everything else on these lists has a '
        'substitute that costs us a day.')

doc.save(OUT)
print('saved', OUT)
