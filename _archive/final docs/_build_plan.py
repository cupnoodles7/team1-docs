#!/usr/bin/env python3
"""Rebuild final_plan.docx (v2) preserving the v1 visual language exactly."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'final_plan_v1_backup.docx'
OUT = 'final_plan.docx'

NAVY   = '13487F'
DEEP   = '0D3562'
INK    = '16233B'
GOLD   = 'C8922F'
TINT   = 'E9F1FB'
KICKER = 'A9C9EE'
SUBTLE = 'D4E3F6'
LABEL  = '8FB4DE'
WHITE  = 'FFFFFF'

doc = Document(SRC)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child is not sectPr:
        body.remove(child)

# ---------------------------------------------------------------- primitives
def para(space_after=None, space_before=None, style=None):
    p = doc.add_paragraph(style=style)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def run(p, text, size=None, bold=None, color=None, font=None):
    r = p.add_run(text)
    if size:  r.font.size = Pt(size)
    if bold:  r.bold = True
    if color: r.font.color.rgb = RGBColor.from_string(color)
    if font:  r.font.name = font
    return r

def rich(p, parts):
    """parts: list of str (plain) or (text, bold) or (text, bold, color)."""
    for part in parts:
        if isinstance(part, str):
            p.add_run(part)
        elif len(part) == 2:
            run(p, part[0], bold=part[1], color=INK if part[1] else None)
        else:
            run(p, part[0], bold=part[1], color=part[2])
    return p

def secnum(n):
    p = para(space_after=0)
    run(p, n, size=9, bold=True, color=GOLD, font='Consolas')

def sectitle(t):
    p = para(space_after=8)
    run(p, t, size=19, bold=True, color=DEEP, font='Georgia')

def h3(t):
    p = para(space_before=14, space_after=4)
    run(p, t, size=13, bold=True, color=NAVY, font='Georgia')

def label(t):
    p = para(space_before=11, space_after=3)
    run(p, t, size=10.5, bold=True, color=INK)

def body_para(parts, space_after=6):
    p = para(space_after=space_after)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p

def bullet(parts):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p

def spacer(pts=6):
    para(space_after=pts)

# ---------------------------------------------------------------- table bits
def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def center_table(t):
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    t._tbl.tblPr.append(jc)

def center_rows(t):
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
        trPr.append(jc)

def set_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(twips)); tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)

def cell_text(cell, parts, size=9, bold=False, color=INK):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if isinstance(parts, str):
        run(p, parts, size=size, bold=bold, color=color)
    else:
        for part in parts:
            if isinstance(part, str):
                run(p, part, size=size, color=color)
            else:
                run(p, part[0], size=size, bold=part[1],
                    color=part[2] if len(part) > 2 else color)

def data_table(headers, rows, widths):
    """Header row navy/white; first data column bold navy."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    center_table(t); center_rows(t)
    for ci, htxt in enumerate(headers):
        c = t.rows[0].cells[ci]
        set_width(c, widths[ci]); shade(c, NAVY)
        cell_text(c, htxt, size=8.5, bold=True, color=WHITE)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            set_width(c, widths[ci])
            cell_text(c, val, size=9, bold=(ci == 0), color=NAVY if ci == 0 else INK)
    spacer(8)
    return t

def callout(kicker, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = None
    center_table(t); center_rows(t)
    c = t.rows[0].cells[0]
    set_width(c, 9792); shade(c, TINT)
    tcPr = c._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side, sz, col in (('top', 4, TINT), ('left', 24, NAVY), ('bottom', 4, TINT), ('right', 4, TINT)):
        e = OxmlElement('w:' + side)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), col)
        b.append(e)
    tcPr.append(b)
    p1 = c.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
    run(p1, kicker, size=8.5, bold=True, color=NAVY)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    run(p2, text, size=9.5)
    spacer(8)

# ---------------------------------------------------------------------- HERO
hero = doc.add_table(rows=1, cols=1)
hero.style = None
c = hero.rows[0].cells[0]
set_width(c, 9792); shade(c, NAVY)

p = c.paragraphs[0]
p.paragraph_format.space_before = Pt(26); p.paragraph_format.space_after = Pt(4)
run(p, 'DELIVERY PLAN  ·  v2  ·  CAP-2 · CAP-3 · CATALOGUE', size=9, bold=True, color=KICKER)

p = c.add_paragraph(); p.paragraph_format.space_after = Pt(6)
run(p, 'team1 — Discovery & Selection', size=28, bold=True, color=WHITE, font='Georgia')

p = c.add_paragraph(); p.paragraph_format.space_after = Pt(20)
run(p, 'Execution model, feature list, four-week build plan, dependency mapping, '
       'work distribution across five, and the ambiguity register.', size=11.5, color=SUBTLE)

for k, v in (
    ('PROGRAMME   ', 'T&F Reader — 8-Week Graduate Prototype'),
    ('TEAM   ', 'Khushi S Shukla · Prayas Yadav · Moktik · Akriti Khetan · Keshav Sharma'),
    ('OWNS   ', 'CAP-2 Institution listing · CAP-3 Institute selection · Catalogue browse & unified search'),
    ('TARGET   ', 'team1 feature-complete on mock data by end of Week 3 · integration closes Week 4'),
    ('WEEKS   ', 'W1 10–14 Aug · W2 17–21 Aug · W3 24–28 Aug · W4 31 Aug–4 Sep 2026'),
):
    p = c.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, k, size=8, bold=True, color=LABEL)
    run(p, v, size=9.5, color=WHITE)

p = c.add_paragraph(); p.paragraph_format.space_after = Pt(18)
spacer(10)

# ============================================================ 01 HOW WE WORK
secnum('01')
sectitle('How we work')

body_para([
    'The design package is finished — eighteen screens, a fixed palette, a typographic scale and a stated '
    'component architecture. That removes the largest source of rework, and it changes what the risk is. '
    'The risk is no longer ',
    ('what should this look like', True), ', it is ',
    ('five people building the same thing five different ways', True),
    '. This section exists because that failure is invisible until Week 3, and by then it is expensive.'
])

h3('The six failure modes this plan is designed against')

data_table(
    ['Failure mode', 'What it looks like in Week 3', 'What prevents it'],
    [
        ['Duplicated components',
         'Three different teals, four Card implementations, merge conflicts in every screen file',
         'One owner for the component library; components merged before the screens that consume them'],
        ['Serialised on the data layer',
         'Four people idle for three days in Week 1 waiting for an adapter',
         'Types and interfaces committed Day 1; stubs land before implementations'],
        ['Access logic leaks into screens',
         'if (tier === "elite") scattered across twelve files; a rewrite when real rules land',
         'resolveAccess is the only source of actions — screens read state, never licence'],
        ['Mock shapes diverge from real',
         'Integration is a rewrite, not a configuration change',
         'Send our assumed contract as a typed shape to wokay and flambeau in Week 1'],
        ['States bolted on last',
         'Loading, empty, error and offline half-work; the demo breaks on a slow network',
         'The state-matrix pass is a scheduled deliverable, not a bug-fix'],
        ['Everyone edits the router',
         'Constant conflicts on navigation and theme files',
         'File-ownership map; shared files require a second reviewer'],
    ],
    [1750, 3900, 4142]
)

h3('The four-phase execution model')

label('Phase 0 — Foundation Sprint · Days 1–2 · all five people, one branch')
body_para('No feature work begins until main contains all of the following. This is two days of apparent '
          'slowness that buys back a week.')
bullet([('Design tokens', True), ' — the colour, typography, spacing and radius tables in the design '
        'specification transcribed verbatim into theme constants. A two-hour job that eliminates every '
        'later colour-drift argument.'])
bullet([('Contracts as types, before implementations', True), ' — ContentItem, Institution, Session, '
        'DataAdapter, and the resolveAccess signature. These are the seams the other four people code against.'])
bullet([('Mock fixtures', True), ' — roughly forty items spanning all three file formats across all three '
        'access tiers, and eight institutions, derived from wokay\'s Week 1 OPDS samples where they exist. '
        'Deliberately includes the awkward cases: missing cover, no ISBN, long title, an Elite item with '
        'canPersist false, an item with NO acquisition link, a zero-result navigation feed, and one error '
        'envelope per code.'])
bullet([('Component inventory and split', True), ' — walk all eighteen screens, list every distinct '
        'component, dedupe, and divide the twenty-two across all five people. The split is in the table below.'])
bullet([('One component skeleton', True), ' — agreed prop-naming, state handling and file layout, so five '
        'authors produce one library rather than five.'])
bullet([('App shell', True), ' — bottom tab navigation with four empty routed screens, merged to main.'])

label('Phase 1 — Components, split five ways, before any feature work · Days 3–5')
body_para([
    'The component library is ', ('built by all five people during the foundation phase', True),
    ', not by one owner during Week 1. Each person takes four to seven components, chosen to be the ones '
    'their own feature will lean on hardest. Nobody starts a feature until the library is merged. This is '
    'what makes feature ownership safe: a feature owner composes from a library that already exists, so '
    'there is never a reason to write a second Card.'
])
body_para('Every component is built against a hidden state gallery route showing all its variants — default, '
          'loading, empty, error, offline, and all three access tiers. Building the states into the '
          'component is the only way they reliably exist later.')

callout('THE ONE RULE THAT MATTERS',
        'A feature may not introduce a new component. If a feature needs something the library does not '
        'have, it is added to the library — in the library\'s files, by whoever needs it, reviewed by the '
        'component\'s original author. Never inside a feature folder. The rule being protected is "one '
        'implementation, one location", not "one author". This is the rule most likely to be quietly '
        'broken under time pressure, and with five authors instead of one it is now easier to break.')

callout('WHAT SPLITTING THE LIBRARY COSTS, AND HOW IT IS PAID FOR',
        'Five authors produce five styles unless something prevents it. Three things do. Design tokens are '
        'written by one person on Day 1, before any component exists, so nobody picks a colour or a spacing '
        'value. Every component is built from the same skeleton — same prop-naming convention, same state '
        'handling, same file layout — agreed in the foundation sprint. And there is a named consistency '
        'review at the end of Day 5, where all twenty-two are looked at side by side in the gallery before '
        'a single feature begins. Skip that review and the drift will not surface until Week 3.')

label('Phase 2 — Feature verticals, owned end to end · Weeks 1–3')
body_para([
    'Each person owns a feature outright — its data access, its logic, its screens and its states — rather '
    'than owning a horizontal layer that everyone else consumes. ',
    ('This is a deliberate change from the previous version', True),
    ', made to reduce hand-offs and the coordination overhead they create. The cost is that the shared '
    'safeguards move earlier in time rather than living in one person, which is what Phase 0 and Phase 1 '
    'now exist to provide.'
])
body_para('Two things stay single-authored regardless: design tokens, and the shared type and adapter '
          'interfaces. If those fragment across five features, a contract change from wokay or flambeau '
          'becomes a five-file edit instead of a one-file edit.')

label('Phase 3 — State-matrix pass · Week 3')
body_para('Every screen is walked against the full matrix — loading, empty-for-query, empty-for-filters, '
          'error, offline, three access tiers, three session states. Screens 14 through 18 in the design '
          'package exist for exactly this purpose. It is scheduled work with an owner, not something that '
          'happens if there is time.')

label('Phase 4 — Integration is an adapter swap · Weeks 3–4')
body_para([
    'If Phases 0 to 3 held, pointing at real wokay and flambeau services is a configuration change and a '
    'week of fixing shape mismatches. ',
    ('If integration turns into a rewrite, the cause will be Phase 0, not Week 3.', True)
])

h3('Cross-cutting practices')
bullet([('Contract-first, not question-first', True), ' — alongside the question lists in Section 09, send '
        'wokay and flambeau the typed shape we are building against and ask them to correct it. Response '
        'rates are far higher against a concrete artefact, and their silence becomes safe: silence means '
        'they accepted our shape.'])
bullet([('Stub-first', True), ' — every external dependency is consumed through an interface we own, with a '
        'stub committed on Day 1. Nothing is ever blocked; it is running on a stub. See Section 08.'])
bullet([('Fixtures are the contract test', True), ' — when a real shape arrives, update the fixture first, '
        'see what breaks, then integrate. Never integrate against a live service as the first contact.'])
bullet([('Daily merge to main', True), ' — no branch lives longer than a day. Long branches plus a shared '
        'component library is the worst combination available to us.'])
bullet([('Dependency board, reviewed daily', True), ' — every dependency sits in exactly one of four states: '
        'Asked, Answered, Stubbed, Integrated. Owned by the lead.'])

h3('Reusable component inventory and the five-way split')
body_para('This is the Phase 0 deliverable, done in advance. Each person takes the components their own '
          'feature leans on hardest, so the author and the heaviest consumer are usually the same person. '
          'All twenty-two are merged before any feature work begins.')

data_table(
    ['Component', 'Author', 'Appears on', 'Notes'],
    [
        ['ContentCard', 'Prayas', '01, 04, 05, 08, 09, 17, 18', 'The most reused component in the app. Renders any content type from one normalised item.'],
        ['SectionHeader', 'Prayas', '01, 06, 08, 09', 'Featured, Browse by Subject, Recently used, All Institutions.'],
        ['SubjectChip', 'Prayas', '01, 09', 'Outlined teal pill.'],
        ['Carousel + PageDots', 'Prayas', '01', 'Featured content.'],
        ['Tabs', 'Prayas', '01, 04, 08, 09', 'Feed tabs are the prominent use; also detail sections.'],
        ['SearchInput', 'Moktik', '01, 06, 08, 09', 'Consumed by both search surfaces — see Section 04.'],
        ['FilterChip', 'Moktik', '09, 12', 'Selected and unselected states.'],
        ['VoiceOverlay', 'Moktik', '11', 'Dark immersive overlay with live transcription.'],
        ['BottomSheet', 'Keshav', '02, 03, 12', '16px top radius, drag handle, 60–70% height. Sign-in, Access Gate and Filter all reuse it.'],
        ['InstitutionRow', 'Keshav', '06, 10', 'Crest, name, country. Reused in Profile for the current selection.'],
        ['TopAppBar', 'Keshav', 'Every screen', 'Navy. Title, back and search-icon variants.'],
        ['ListRow', 'Keshav', '10', 'Settings and profile rows.'],
        ['AccessTierBadge', 'Akriti', '01, 04, 05, 08, 09, 18', 'Open Access green, Subscription blue, Elite purple. Driven by the resolver, never by raw licence data.'],
        ['ActionButton', 'Akriti', 'Every screen', 'The most variants of anything here: Read, Download, Borrow, Sign in, Waitlist, disabled. Subscribe is deleted with the individual-subscriber scope.'],
        ['ActionBar', 'Akriti', '04, 05', 'Sticky bottom bar. Renders purely from the actions array.'],
        ['Skeleton (card / block)', 'Khushi', '01, 04, 05, 06, 08, 09, 15, 16', 'No spinners anywhere. Dimensions must match final content to avoid layout shift.'],
        ['EmptyState', 'Khushi', '06, 08, 09, 17', 'Two copy variants: no results for a query, no results for filters.'],
        ['ErrorState', 'Khushi', '04, 05, 14', 'Retry and learn-more actions. No stack traces surfaced.'],
        ['OfflineBanner', 'Khushi', 'Global, 15', 'Persistent dark grey. Library and institution list stay usable behind it.'],
        ['BottomTabBar', 'Khushi', '01, 08, 09, 10', 'Catalogue, Search, Library, Profile.'],
        ['FormatSelector', 'Khushi', '05', 'AT RISK OF DELETION — wokay confirmed 11 Aug that one work has exactly one format, so there is nothing to select between. Decide at Wednesday standup: drop it, or keep it as a single-format display strip.'],
        ['ProgressBar', 'Khushi', '07, 08', 'Download and reading progress. Consumed by t4targaryen.'],
    ],
    [1750, 900, 2050, 5092]
)
body_para([
    ('Design tokens are not in this table on purpose.', True),
    ' They are written by one person on Day 1, before any component exists, so that no component author '
    'ever picks a colour or a spacing value. That single exception is what makes splitting the rest safe.'
])

# ==================================================== 02 SCOPE AND USER FLOW
secnum('02')
sectitle('Scope and user flow')

body_para([
    'The catalogue is ', ('fully visible before sign-in', True),
    ' and ', ('scoped to entitlements after it', True),
    '. A logged-out user sees every title from every institution and can consume anything Open Access. '
    'Signing in narrows the feed to what that user is actually entitled to, plus Open Access. Access is '
    'resolved per item, and every item surfaces exactly one action.'
])

h3('The two states')
body_para([('There were three. The individual-subscriber state is deleted', True),
           ' — wokay cut individual subscribers at their own gate, and their entitlement resolver is keyed '
           'by institution. What replaces it is not a third session type but an ANONYMOUS surface, which is '
           'narrower than the old logged-out state: the unauthenticated feed serves open access only, not '
           'the full catalogue.'])
data_table(
    ['State', 'Feed shows', 'What they can do', 'Everything else'],
    [
        ['Anonymous', 'Open access ONLY, from GET /opds/v1/public/catalogue. NOT the full catalogue — the full catalogue is entitlement-scoped and token-protected',
         'Read and Download, in whichever formats the assets declare. No sign-in, no loan, no key',
         'Everything else is simply absent from this feed rather than gated in it. The Access Gate appears when a user arrives at a gated item through a link or through pending intent'],
        ['Signed in — institutional', 'The institution\'s root feed, already scoped by the backend to what it is entitled to, plus open access, which appears for every institution regardless of entitlement',
         'Open access: Read and Download. Subscription: Borrow, then Read and Download. Elite: Borrow, then Read only — and Waitlist when no copies are free',
         'A non-entitled title is reachable only through pending intent, and renders as Access Restricted with no action. So is an item sent with no acquisition link'],
    ],
    [1300, 2900, 3300, 2292]
)

h3('Journey')
bullet([('App opens.', True), ' A dismissible sign-in sheet appears. The user can sign in now, or close it '
        'and browse. Screen 02.'])
bullet([('Browse.', True), ' Closing the sheet reveals the full catalogue — every institution, all four '
        'content types, no wall. Screen 01.'])
bullet([('Open Access is fully usable with no session', True), ' — Read, and Download in whichever formats '
        'the OPDS feed declares. No loan is created, because there is no borrower yet.'])
bullet([('Tapping a Subscription or Elite title raises the Access Gate.', True), ' The tapped item is '
        'written to the pending-intent store ', ('before', True), ' the gate opens. Screen 03.'])
bullet([('Two routes.', True), ' Through my institution, or browse free content. The second was '
        '"personal account" — individual subscribers are cut from scope, so it becomes the anonymous '
        'open-access feed instead.'])
bullet([('Institution route', True), ' — searchable institution list, select, then hand off to flambeau '
        'for SAML with signIn.idpHint attached. Sign-in is ALWAYS SAML, so there is no route to choose '
        'between. ', ('Free-content route', True),
        ' — wokay\'s public feed, no token. Screen 06; screen 13, email sign-in, has no owner under this '
        'model.'])
bullet([('Return to the same item.', True), ' On re-entry, pending intent restores the exact title the user '
        'tapped and re-resolves its access against the new session. This is the step most likely to be cut '
        'under pressure and the one users will notice most.'])
bullet([('From then on', True), ' the feed re-scopes — and wokay do the scoping. The signed-in root feed '
        'is already filtered to what the institution is entitled to, plus open access, which appears for '
        'every institution regardless. We do not duplicate their entitlement logic.'])

h3('Runtime ownership — who does what, in order')
body_para([
    'Summarised from the outside, this product sounds like it belongs to the other three teams: the user '
    'signs in with flambeau, wokay serve the catalogue, t4targaryen render the reader. That summary '
    'describes a user who already knows what they want and is already signed in — which is a deep link, '
    'not an application. ',
    ('The other teams own capabilities. team1 owns the application that composes them.', True)
])
data_table(
    ['#', 'What happens at runtime', 'Owner'],
    [
        ['1', 'App launches; ONE root feed fetched, its navigation rows read, normalised to one model, rendered as shelf tabs', 'team1 — wokay serve the bytes and decide the shelves'],
        ['2', 'User browses, filters, sorts, paginates by following next links', 'team1 — the query surface. wokay run the search itself, server-side and entitlement-scoped'],
        ['3', 'User taps a gated title. Access resolves to requires_signin, the item is written to pending intent, the Access Gate is raised', 'team1 — neither other team knows this happened'],
        ['4', 'User chooses "through my institution": list, search, select, persist', 'team1 — OPDS does not model institutions'],
        ['5', 'Handoff with institution context; SAML or SSO authenticates, possibly outside the app', 'flambeau — team1 is the caller'],
        ['6', 'Return caught, pending intent restored, session read, feed re-scoped', 'team1'],
        ['7', 'Restored item re-resolved: the acquisition link interpreted against the session and the loan state', 'team1 — this is where canPersist false becomes a missing Download button'],
        ['8', 'User taps Read; handoff with item and entitlement context', 't4targaryen — team1 is the caller'],
        ['9', 'Reader renders the content', 't4targaryen'],
    ],
    [500, 6100, 3192]
)
callout('WHY THIS MATTERS BEYOND SCOPE DEFENCE',
        'wokay know what an item is. flambeau know who the user is. Neither knows what this user may do '
        'with this item, and that join only exists where both facts meet — step 7, on the client. There is '
        'no step in this sequence at which another team could compute it without being handed the other '
        'team\'s data. Fifteen of the eighteen designed screens are ours; wokay render nothing, flambeau '
        'render one, t4targaryen render two.')

h3('The action matrix')
body_para('One state produces one action. This table is the specification for resolveAccess and the only '
          'place action logic is allowed to live.')

callout('CORRECTED AGAINST WOKAY\'S SOURCE OF TRUTH — READ THIS BEFORE THE TABLE',
        'The previous version of this table was inverted on two of the three tiers. It derived from the '
        'rule "a loan is required for anything the user can take a copy of", which is not the rule the '
        'backend implements. wokay\'s section 02 and flows A, B and C state the actual rule, and it changes '
        'the action on Open Access and on Elite. Design Specification section 3.1 carried the same error '
        'and has been corrected to v0.3, which is an edit to a signed document and needs ratification. '
        'Anyone who saw the earlier matrix should be told explicitly that it was wrong.')

data_table(
    ['Tier', 'Logged out', 'Signed in, entitled', 'Signed in, not entitled'],
    [
        ['OPEN_ACCESS', 'Read + Download', 'Read + Download — identical. No entitlement, no licence, no loan record, ever',
         'Read + Download — open access items appear for every institution regardless of entitlement'],
        ['SUBSCRIPTION', 'Sign in', 'Borrow, then Read + Download. Encrypted, loan-backed, but UNLIMITED — nothing to reserve, so no queue can ever form', 'Access Restricted, no action'],
        ['ELITE', 'Sign in', 'Borrow, then Read ONLY. Download refused server-side (canPersist false)', 'Access Restricted, no action'],
        ['ELITE, no copies free', '—', 'Waitlist — writes a hold, returns a queue position. The ONLY tier a queue can form on', '—'],
    ],
    [1600, 1800, 3400, 2992]
)

callout('THE LOAN RULE, CORRECTED, IN ONE LINE',
        'A loan exists for anything ENCRYPTED — Subscription and Elite. A lease and a queue exist only '
        'where copies are FINITE, which is Elite alone. Open access is stored as plaintext precisely so an '
        'anonymous reader can open it, so it has no entitlement, no licence and no loan record in any '
        'session state. Subscription is encrypted and loan-backed but unlimited, so it is downloadable and '
        'can never queue. Elite is the only tier that borrows against a finite pool, and it is the only '
        'tier that cannot be downloaded.')

callout('THE TWO CORRECTIONS, STATED PLAINLY',
        'Open Access does NOT require a Borrow once signed in. Signing in changes nothing about it. And '
        'Elite is NOT read-only-without-a-loan — it is the ONE tier that borrows, consumes a copy and can '
        'queue. The previous version had these exactly backwards, which would have produced a resolver '
        'that was wrong on two of three tiers and a demo that showed the wrong button on most of the '
        'catalogue.')

callout('THIS IS NOT THE "BORROW" THAT WAS REMOVED',
        'An earlier version of this plan removed a Borrow action, and reviewers who saw it will misread '
        'this one. That Borrow was a purchase-adjacent path for individual users on content they were not '
        'entitled to, and it stays removed along with Buy. This Borrow is loan creation on content the user '
        'is already entitled to — the library lending model, flambeau\'s CAP-4. Different trigger, '
        'different owner, different meaning. Say so when circulating.')

body_para([('Three consequences worth stating plainly.', True),
           ' Subscribe is GONE, not B2C-only — individual subscribers are cut from scope by wokay\'s gate '
           'decision 7, and the entitlement resolver is keyed by institution. no_seats is live but narrow: '
           'Elite only, because it is the only tier with a finite copy count. And no_seats is unreachable '
           'on a catalogue card, because the feed carries copies.total and never copies.available — the '
           'app fetches availability per item on the detail screen, where the decision is actually being '
           'made. FL-11\'s contingency is therefore the design, not the fallback.'])

callout('THE RULE THAT REPLACES TIER-BASED RESOLUTION',
        'The acquisition link decides the buttons, not our own logic. wokay send an OPDS acquisition link on '
        'every publication, and that link is the input: rel=open-access means Read and Download; '
        'rel=acquisition with canPersist true means Borrow then Read and Download; rel=borrow with '
        'canPersist false means Borrow then Read only; and if they do not send a link, the button does not '
        'exist. resolveAccess becomes a link-and-properties INTERPRETER, not a rules engine. No branch in it '
        'may read accessTier to decide an action — the tier survives as the badge label. wokay ask for this '
        'in writing by end of Week 1, and if we stay silent they document it as a contract test on their '
        'side. Design Specification 5.1 — "the UI must never calculate access rights" — is the same rule '
        'stated at a different altitude.')

callout('DESIGN PRINCIPLE',
        'One item model, one access-resolution layer, two feed scopes. The card, the badge and the detail '
        'view never change shape — only the scope of the feed and the single resolved action change between '
        'anonymous and signed in. Everything downstream of resolveAccess is presentation.')

callout('CHANGED FROM v1 — AND THE DESIGN PACKAGE NEEDS THE SAME EDITS',
        'Five reversals. (1) The signed-in feed is scoped by the backend, and the anonymous feed is '
        'OPEN ACCESS ONLY — not the full catalogue. v1 and Design Specification 4.1 both state that the '
        'catalogue is "never filtered by institution", as does step 2 of the flow in index.html. Those have '
        'been corrected. L-2 has changed character: wokay CANNOT SERVE the v1 model, so this is now a '
        'statement of backend fact requiring ratification rather than a request. (2) Buy and Borrow are '
        'removed; SUBSCRIBE IS ALSO REMOVED, along with the individual-subscriber scope entirely — wokay '
        'gate decision 7. (3) The plan runs to four weeks, with team1 feature-complete at the end of Week 3. '
        '(4) Search is re-scoped DOWNWARD, not upward: catalogue search is server-side and entitlement-'
        'scoped, so matching, tokenisation and ranking are wokay\'s, not ours. This reverses the Section 03 '
        'audit conclusion. (5) There are not three OPDS feeds. There is ONE root feed per institution '
        'carrying navigation rows and groups, plus one public open-access feed. Shelves are discovered and '
        'configured per institution, so the tab bar must be data rather than a hardcoded list of three.')

callout('THE EDGE CASE THIS CREATES',
        'Scoping the feed after sign-in means pending intent can return a user to an item that is no longer '
        'in their feed. With Subscribe deleted, there is exactly one resolution: it can only happen on a '
        'Subscription or Elite title the institution does not hold — open access is always in every feed — '
        'and that item renders as not_entitled with no action: screen 14, Access Restricted. There is now a '
        'SECOND path to the same state, which is an item wokay send with no acquisition link at all. That is '
        'why D8 and screen 14 survive despite the feed itself never listing non-entitled content, and why an '
        'empty actions array is a normal input to ActionBar rather than an edge case.')

# ============================================================== 03 FEATURES
secnum('03')
sectitle('Features')

h3('A · Catalogue browse')
body_para([
    ('CORRECTED. There are not three OPDS endpoints.', True),
    ' wokay serve ', ('one root catalogue feed per institution', True),
    ' — GET /opds/v1/institutions/{id}/catalogue — carrying navigation rows and groups, which are shelves. '
    'One request draws the whole home screen. A shelf is then paged via GET .../groups/{gid}. Separately '
    'there is one public open-access feed, GET /opds/v1/public/catalogue, needing no token. Shelves are '
    'therefore discovered per institution and CONFIGURED per institution — wokay\'s admin console has a '
    'per-institution catalogue config setting feed title, shelf order and page size — so nothing may '
    'hardcode a list of three.'
])
callout('WHAT THIS CHANGES, AND WHAT IT DOES NOT',
        'The tabs decision survives and is strengthened: the tab bar MUST be driven from data, because the '
        'shelf set is the institution\'s to configure and can change without a rebuild. What disappears is '
        'the premise that the three feeds are split by content type. wokay\'s contentType is PDF, EPUB or '
        'AUDIO — a FILE FORMAT, offered as a free query parameter — not a work type, and it does not '
        'partition the feed. The cross-feed deduplication and merged-pagination problems also disappear, '
        'because there is one feed with groups rather than three independent cursors. L-5 stays open on the '
        'narrower question of whether shelves present as tabs at all.')
data_table(['Ref', 'Deliverable'], [
    ['A0', 'Shelf tabs on catalogue and search, read from the root feed\'s navigation rows. Data, never a hardcoded list — the shelf set is per-institution configuration  — CHANGED'],
    ['A1', 'Anonymous browse via GET /opds/v1/public/catalogue. OPEN ACCESS ONLY, not the full catalogue — the full catalogue is entitlement-scoped and token-protected  — CHANGED, and it is a reduction'],
    ['A2', 'Signed-in feed, scoped by the BACKEND to what the institution is entitled to, plus open access, which appears for every institution regardless of entitlement  — CHANGED FROM v1'],
    ['A3', 'Item card carrying title, publisher, file format and access-tier badge'],
    ['A4', 'Pagination — follow the next link until it is absent, never count pages. page/size/total exists ONLY on the institutions REST endpoint. Two models split cleanly by endpoint family, so the "one interface over both" abstraction is dropped  — SHRUNK'],
    ['A5', 'Featured carousel and Browse-by-Subject chips (screen 01). Subject chips need OPDS facets, which wokay build ONLY IF WE ASK — see Section 09'],
    ['A6', 'Access Gate raised when a locked item is opened while anonymous. Second option becomes "browse free content", not "personal account"  — CHANGED'],
    ['A7', 'Feed re-scopes and re-resolves on sign-in and sign-out without a full reload'],
    ['A8', 'Follow hrefs, never build URLs. catalogueUrl is handed over on institution detail; every later hop is an href from a feed. The single permitted construction is expanding ".../search{?query}"  — NEW, and it is a hard rule'],
    ['A9', 'Zero results is a navigation feed, not an empty array. OPDS forbids empty arrays, so a zero-result search returns a feed with a navigation entry pointing back at the catalogue. Render a browse-instead affordance; a missing publications key is NOT an error  — NEW'],
], [700, 9092])

h3('B · Search and filter')
callout('THE SECTION 03 AUDIT CONCLUSION IS REVERSED. READ THIS FIRST',
        'The previous version audited OPDS 2.0, found it specifies a feed format rather than a search '
        'engine, and concluded that matching, tokenisation and ranking are therefore ours — making this the '
        'largest block in the plan. That reasoning was sound about the SPECIFICATION and wrong about the '
        'PROJECT, because it audited the format instead of asking the team who serve it. wokay: "Catalogue '
        'search has to be server-side, because results are filtered by what the institution is entitled to. '
        'Searching on the client would mean shipping the whole catalogue to the phone and reimplementing '
        'entitlement rules there, and a member must never see a result they cannot open. So: we filter, you '
        'render." Matching, tokenisation and ranking are all THREE theirs. This is the single largest scope '
        'reduction in the plan and it lands on the most-loaded person.')

body_para([
    'Two of the three searches in this project are server-side, and only the third is ours. ',
    ('Catalogue search is a wokay endpoint with a team1 UI. Institution search is a wokay endpoint with a '
     'team1 UI. In-file search is built by wokay and queried by t4targaryen.', True),
    ' team1 owns no search engine at all — we own two query surfaces and the state behind them.'
])

data_table(
    ['Requirement', 'Who actually does it', 'What team1 builds'],
    [
        ['Search — title, author, subject, description', 'wokay, server-side, entitlement-scoped. METADATA ONLY — it does not search inside books', 'Query state, expanding the templated search link, rendering, paging by next link'],
        ['DOI / ISBN exact-match', 'CORRECTED 11 Aug. The top-level isbn in their document does not exist in the feed — it arrives wrapped in a longer identifier string, and titles without an ISBN carry an internal identifier instead. There is NO DOI, and there never will be', 'URN extraction is back, and it is a real job again. DOI comes OFF screen 04 — a design change, not an open ask'],
        ['Filter — content type (PDF, EPUB, AUDIO)', 'wokay, free. A fixed enum, a query parameter on endpoints we already call', 'Three hardcoded chips that send the parameter'],
        ['Filter — access tier', 'REOPENED 11 Aug. There is no tier field — we derive the tier from the licence model. So it is unclear whether the tier can still be sent as a filter parameter, or whether they map it server-side', 'The chips, but ASK FIRST. Moktik builds them Day 3, so this needs a same-day answer'],
        ['Filter — subject', 'wokay, BUT ONLY IF WE ASK. Subjects are dynamic, so they cannot be known per institution client-side. Needs OPDS facets, roughly half a day of theirs', 'The chips — and the Week 1 answer that causes the facets to exist at all'],
        ['Filter — publisher, year', 'wokay, unconfirmed that the fields are consistently populated', 'The chips, hiding any dimension with no data'],
        ['Filtering happens BEFORE pagination', 'wokay — filters are query parameters, deliberately', 'Never filter inside a fetched page. "Filter within a page of twenty and you show three results while the pager still says page 1 of 12"'],
        ['Visible result count', 'Unconfirmed on OPDS feeds. total IS returned by the institutions REST endpoint', 'Count locally and approximate on catalogue surfaces'],
        ['Sort options', 'Unconfirmed. sort=publishedAt,desc is HINTED on the groups endpoint but not stated as the contract', 'Ask; sort locally over the fetched page meanwhile'],
        ['Combined search + filter + sort surviving navigation', 'Nothing from anyone', 'Everything. This is genuinely ours and is what remains of the "largest block" claim'],
        ['The same search on the institution list', 'wokay, server-side and paged: ?q=&country=&page=&size=', 'The UI, and recently-used pinning, which is client-side'],
        ['Distinct empty states', 'Nothing — but a zero-result response is a NAVIGATION FEED, not an empty array', 'Three variants, including browse-instead. A missing publications key is not an error'],
        ['Voice search', 'Nothing', 'React Native speech-to-text producing a string into the same pipeline'],
    ],
    [2100, 4000, 3692]
)

callout('TWO RIDERS THAT BELONG IN THE UI, NOT ONLY IN THE CODE',
        'Catalogue search is METADATA ONLY — title, authors, subjects and description. It does not search '
        'the text inside books. That must be reflected in the placeholder or helper copy, or a reviewer '
        'typing a phrase from page 88 will file a bug. And the search endpoint lands in WEEK 4, which is '
        'team1\'s integration and BDD week — so the search pipeline is built against fixtures for three '
        'weeks and integrated in the busiest week of the plan. That is a real schedule risk and it should be '
        'on the dependency board in Week 1 rather than discovered in Week 4. The mitigation is the interface '
        'boundary: fetch-and-filter is a fixture stand-in only and must never become the design.')

data_table(['Ref', 'Deliverable'], [
    ['B1', 'Catalogue search UI over wokay\'s server-side endpoint. Query state, templated-link expansion, results rendering, paging by next link, empty and error states. Matching, tokenisation and ranking are THEIRS  — CHANGED, and it is a large reduction'],
    ['B2', 'ISBN unpacked from the identifier string — URN extraction IS required, reversing the previous note. Titles with no ISBN carry an internal identifier. DOI does not exist and will not: it comes off screen 04  — CHANGED AGAIN, 11 Aug'],
    ['B3', 'Filter — file format, PDF / EPUB / AUDIO. A free fixed-enum query parameter. It does NOT partition the feed, so it no longer collapses into the tabs  — CHANGED'],
    ['B4', 'Filter — subject, publisher and publication year. Subject needs OPDS facets, which exist only if we ask for them in Week 1'],
    ['B5', 'Filter — access tier. UNBLOCKED: accessTier already exists on wokay\'s record and is offered as a free filter  — CHANGED, this was the hardest dependency in the plan'],
    ['B6', 'Sort options — a "Sort by" facet group or a query parameter, agreed with wokay. Still open'],
    ['B7', 'Visible result count. Approximate locally on OPDS surfaces; total is real on the institutions endpoint'],
    ['B8', 'Combined search, filter and sort state that survives navigation — entirely ours'],
    ['B9', 'Institution search UI over wokay\'s paged server endpoint ?q=&country=&page=&size=. A separate pipeline sharing only the shell  — CHANGED from client-side'],
    ['B10', 'Three empty states: no-results-for-query, no-results-for-filters, and browse-instead from a zero-result navigation feed  — CHANGED'],
    ['B11', 'Voice search on the catalogue only — speech-to-text producing a string into the B1 pipeline'],
    ['B12', 'Recent queries and popular subjects on the search landing surface (screen 09)'],
    ['B13', 'Filter and sort bottom sheet (screen 12)'],
], [700, 9092])

h3('C · Institution listing and selection   (CAP-2 + CAP-3)')
data_table(['Ref', 'Deliverable'], [
    ['C1', 'Browse and search available institutions, with recently-used pinned above the full list'],
    ['C2', 'Institution detail view'],
    ['C3', 'Select an institution to scope catalogue and access'],
    ['C4', 'Selection persisted and restored on relaunch'],
    ['C5', 'Change or clear the selected institution, from Profile'],
    ['C6', 'Route to sign-in by passing signIn.idpHint to flambeau. Institutional sign-in is ALWAYS SAML — there is no authMethod field on wokay\'s institution record and no branch to write. authRouting stops being a branching file and becomes a pass-through  — CHANGED'],
    ['C7', 'Entry choice: find your institution, or browse free content. Individual subscribers are cut from scope (wokay gate decision 7), so the second option is the ANONYMOUS PUBLIC FEED rather than a personal-account path. Screen 03 still belongs to the access layer  — CHANGED'],
    ['C8', 'Institution list works offline from cache  — NEW. It sits on the authentication critical path; if it fails, the user cannot sign in at all'],
    ['C9', 'An inactive institution returns 404, not 403, deliberately, so its existence is not disclosed. Treat 404 on institution detail as not-found and never as forbidden  — NEW'],
], [700, 9092])

callout('THE EARLIEST REAL THING TEAM1 GETS FROM ANYONE',
        'The institution endpoints are WEEK 2 — ahead of the root feed in Week 3, and ahead of search, '
        'shelves and the public feed in Week 4. GET /api/v1/institutions with q, country, page and size, '
        'returning items, page, size and total, no token. GET /api/v1/institutions/{id} returning branding, '
        'signIn.idpHint and catalogueUrl, no token. Both schemas are published in full, so W-3 and W-17 are '
        'answered and the fixtures can be written against their exact field names. Keshav is the first person '
        'on the team who can delete a mock.')

h3('D · Access resolution and actions')
body_para('The core of this build. One resolution layer decides an item\'s state; the state decides the '
          'single action.')
data_table(['Ref', 'Deliverable'], [
    ['D1', 'resolveAccess(item, session, loan, availability) producing tier, state and actions BY INTERPRETING THE ACQUISITION LINK. A link-and-properties interpreter, not a rules engine. No branch may read accessTier to decide an action  — CHANGED, and the signature changed with it'],
    ['D2', 'Action derived from the link\'s rel and canPersist, so every item surfaces exactly one correct action  — CHANGED'],
    ['D3', 'Access-tier badge on every catalogue card and every detail view. This is now the ONLY consumer of accessTier'],
    ['D4', 'Elite offers no Download. Enforced by canPersist false on the link, and refused server-side with DOWNLOAD_NOT_PERMITTED regardless — so the rule lives on the server, not only in our UI'],
    ['D5', 'REMOVED: the Subscribe action, together with the individual-subscriber scope. wokay gate decision 7 — the entitlement resolver is keyed by institution, and the anonymous public feed is the skip-institution path instead  — CHANGED FROM v1'],
    ['D6', 'Read handoff to t4targaryen'],
    ['D7', 'Download handoff on Open Access and Subscription only, with formats read from the assets declaration'],
    ['D8', 'not_entitled state rendered (screen 14). Two paths now: pending intent onto a non-entitled title, and an item sent with no acquisition link at all  — CHANGED'],
    ['D9', 'Borrow creating a loan through flambeau CAP-4, on entitled SUBSCRIPTION and ELITE content. NEVER on open access, which has no loan record in any session state  — CORRECTED, this was backwards'],
    ['D10', 'Loan state read from flambeau and joined to each item, so a borrowed title shows Read rather than Borrow  — NEW'],
    ['D11', 'Optimistic transition — the button flips from Borrow on success, without refetching the feed. On Elite it is two-outcome: 201 with a loan, or 202 with a queue position'],
    ['D12', 'no_seats and waitlist — ELITE ONLY, and DETAIL SCREEN ONLY. The feed carries copies.total and never copies.available, so a card can never render it; availability is a separate per-item flambeau call. FL-11\'s contingency is the design  — CORRECTED AND NARROWED'],
    ['D13', 'Borrow failure handling driven by wokay\'s enumerated error codes — NO_ENTITLEMENT, ENTITLEMENT_EXPIRED, ENTITLEMENT_SUSPENDED, CONTENT_NOT_READY, DOWNLOAD_NOT_PERMITTED, FORBIDDEN_INSTITUTION_MISMATCH  — CHANGED: there is now a published error model'],
    ['D14', 'One error-copy map keyed on the error CODE, not on HTTP status. Denials carry a reason so the app can say "your library\'s subscription has expired" rather than "not available". NOT_FOUND must stay deliberately vague — wokay return the same 404 for unknown, archived and not-entitled so the catalogue cannot be mapped by walking ids  — NEW'],
    ['—', 'REMOVED: Buy; the earlier purchase-adjacent Borrow; and Subscribe. See the note in Section 02'],
], [700, 9092])

h3('E · Item detail')
data_table(['Ref', 'Deliverable'], [
    ['E1', 'Book detail — cover, format selector, ISBN (a top-level field, no URN parsing), publisher, subjects (screen 05)'],
    ['E2', 'Article detail — journal metadata, abstract, tabs (screen 04). DOI is OMITTED until wokay add the field; there is no DOI anywhere in their schema  — CHANGED'],
    ['E3', 'Journal and audio detail rendered through the same normalised model'],
    ['E4', 'Identical access-state and action model to the catalogue list, PLUS the two things only detail can do: the availability call and therefore the no_seats state  — CHANGED'],
    ['E5', 'Screens 04 and 05 differ by WORK TYPE — book, journal, article, audiobook — which wokay do not model at all. Their contentType is a file format. Until they add it, the work type is DERIVED, isolated in one function, and both screens run on derived data  — NEW, and it is the Day 1 ask'],
], [700, 9092])

h3('F · Platform and shell')
data_table(['Ref', 'Deliverable'], [
    ['F1', 'OPDS 2.0 consumption normalised to one internal content model. ONE adapter over one root feed and its groups, not three adapters over three endpoints  — CHANGED'],
    ['F2', 'Data access layer with swappable mock and real adapters, so integration is a configuration change. Every catalogue method takes an HREF we were given, never an id we templated — the signature is what stops URL construction creeping back in  — CHANGED'],
    ['F9', 'items:batch — POST /api/v1/catalogue/items:batch, capped at 100 ids, returning items, notFound and denied as three separate keys so one bad id does not break a screen and the client can tell gone from not-yours. Any surface holding item ids without metadata needs this  — NEW'],
    ['F10', 'A typed error layer over wokay\'s envelope — timestamp, status, code, message, path, traceId — with the code enumerated. This is what D14\'s copy map consumes  — NEW'],
    ['F3', 'Session and entitlement context consumed from flambeau and applied to feed scoping and resolution'],
    ['F4', 'App shell, bottom tab navigation, design tokens and the shared component library'],
    ['F5', 'Loading, empty, error and offline states across every surface'],
    ['F6', 'Pending-intent store — the tapped item survives the entire authentication round trip  — NEW'],
    ['F7', 'Profile screen, institution-management section only (screen 10)'],
    ['F8', 'Loan state cached client-side with invalidation on borrow, expiry and sign-out, and a defined offline behaviour  — NEW'],
], [700, 9092])

h3('G · Integration')
body_para([('Replaces the v1 "P2 — after Week 2" section, which is removed.', True),
           ' None of this is new feature work; it is swapping stubs for real services as they land.'])
data_table(['Ref', 'Deliverable'], [
    ['G1', 'Swap the mock adapter for real wokay OPDS catalogue and institution feeds'],
    ['G2', 'Wire real flambeau session, entitlements and authentication routing'],
    ['G3', 'Wire the real reader launch to t4targaryen'],
    ['G4', 'Accessibility pass and BDD test suite, required for the Week 5 Integration Readiness Gate'],
], [700, 9092])

h3('Screen ownership across the design package')
body_para('Eighteen screens were designed; team1 does not build all of them. This table settles the boundary '
          'before anyone starts.')
callout('WHICH DESIGN ARTEFACT IS AUTHORITATIVE',
        'The eighteen PNGs were generated against v0.1 and four of them are now wrong: 01 shows a single '
        'mixed catalogue with no feed tabs, 03 does not name the item pending intent will return to, 09 is '
        'not scoped per feed, and 18 predates the current action matrix. Those four are superseded by live '
        'HTML mockups in the Live Mockups v0.2 section of index.html, which render in a browser and can be '
        'edited directly. Screens 02, 04–08 and 10–17 are unchanged and their PNGs remain accurate. Build '
        'against the mockups where the two disagree.')
data_table(
    ['Screen', 'Built by', 'Owner / note'],
    [
        ['01 Catalogue Home', 'team1', 'Prayas — including the feed tabs (A0)'],
        ['02 Sign-In Sheet', 'team1 renders, flambeau authenticates', 'Keshav'],
        ['03 Access Gate', 'team1', 'Akriti — sheet and routing decision, over Keshav\'s BottomSheet'],
        ['04 Article Detail', 'team1', 'Khushi'],
        ['05 Book Detail', 'team1', 'Khushi'],
        ['06 Institution List', 'team1', 'Keshav — CAP-2'],
        ['07 Reader View', 't4targaryen', 'We launch it; we do not build it'],
        ['08 Library / Downloads', 't4targaryen', 'Including the guest-downloads section'],
        ['09 Search & Filter', 'team1', 'Moktik — search is scoped to the active feed tab'],
        ['10 Profile', 'Split', 'Keshav — institution-management section only, see Q-K'],
        ['11 Voice Search', 'team1', 'Moktik'],
        ['12 Filter & Sort Sheet', 'team1', 'Moktik — over Keshav\'s BottomSheet component'],
        ['13 Email Sign-In', 'flambeau', 'We route to it'],
        ['14 Error States', 'team1', 'Khushi — she renders the state, Akriti decides when it applies'],
        ['15 Offline Banner', 'team1', 'Khushi'],
        ['16 Loading Skeleton', 'team1', 'Khushi'],
        ['17 Empty States', 'team1', 'Khushi renders; Moktik supplies which of the two variants applies'],
        ['18 Access Tier Variants', 'team1', 'Akriti — the action-states reference for the whole team'],
    ],
    [2100, 3300, 4392]
)

# ================================================== 04 SEARCH ARCHITECTURE
secnum('04')
sectitle('Search architecture — catalogue and institutions')

body_para([
    ('The decision: build the shell once, the pipelines twice, the result contract once.', True),
    ' Roughly seventy per cent shared, thirty per cent deliberately duplicated. This reverses the v1 '
    '"one engine, two surfaces" position, which was argued from a two-week deadline. With time removed as '
    'a constraint, the deadline argument disappears and the engineering argument points the other way.'
])

body_para([
    ('The OPDS audit settles this rather than complicating it.', True),
    ' OPDS 2.0 does not model institutions at all — they are not publications, so there is no feed, no '
    'search link, no facet group and no shared backend contract available to us. Catalogue search talks to '
    'a wokay search endpoint; institution search talks to a cached list we hold ourselves. ',
    ('Below the user interface there is simply nothing to share.', True)
])

body_para([
    'One further consequence of the three-feed decision: ',
    ('catalogue search is scoped to the active feed tab', True),
    ', so it queries one endpoint and holds one cursor at a time. This avoids having to re-rank a union of '
    'three independently-ranked result sets against a scoring function of our own — a problem we would '
    'otherwise have inherited on top of building the ranking itself. The shared shell is unaffected; only '
    'the catalogue pipeline gains a feed parameter.'
])

h3('Why they are not the same problem')
data_table(
    ['Dimension', 'Catalogue search', 'Institution search'],
    [
        ['OPDS support', 'A search link and template only — no matching, no ranking, no sort', 'None. Institutions are not publications'],
        ['Corpus', 'Tens of thousands, paginated', 'Tens to low thousands, bounded'],
        ['Execution', 'Server-side request, but we own matching and ranking', 'Client-side over a cached list — entirely ours'],
        ['Fields', 'title, author, subject, and DOI/ISBN via URN extraction', 'name, aliases, country, possibly email domain'],
        ['Matching', 'Fuzzy, with a separate exact-match path for DOI and ISBN', 'Fuzzy and prefix, with diacritic folding — Zürich must match Zurich'],
        ['Filters', 'Six dimensions, five of which we build outright', 'None, or country alone'],
        ['Ranking', 'Ours — OPDS supplies no relevance', 'Recently-used pinned above everything, then alphabetical'],
        ['Voice input', 'Yes', 'No — explicitly out of scope'],
        ['Offline', 'Degrades to the cached page', 'Must work fully offline'],
        ['Failure mode', 'An empty result list', 'A user who cannot sign in, plus a "can\'t find your institution?" escape hatch'],
    ],
    [1500, 3400, 4892]
)

h3('The three layers')
label('Shared — the search shell')
body_para('Query input, debounce, the query state machine (idle → typing → searching → results → empty → '
          'error), recent queries, list virtualisation, result count, clear affordance, and the loading, '
          'empty and error components. This is the bulk of the code and it is genuinely identical.')

label('Separate — the two query pipelines')
body_para('Each owns its own fetching, matching, ranking and filtering. They share no code and can change '
          'independently.')

label('Shared — the result contract')
body_para('Both pipelines return the same shape, so the shell never knows which one it is talking to.')

h3('The argument that decides it')
body_para([
    ('Institution search is on the authentication critical path. Catalogue search is not.', True),
    ' If catalogue search breaks, a user browses instead. If institution search breaks, the user cannot '
    'sign in, cannot obtain entitlements and cannot read anything. The two have different reliability '
    'requirements, different offline requirements and different owners. Coupling them means a catalogue '
    'search regression can take down sign-in — and that is not a trade worth making at any deadline.'
])
body_para([
    'The secondary argument is what a single configurable engine actually becomes: an options object with '
    'fifteen flags, where each flag is used by exactly one caller. ',
    ('That is not reuse, it is coupling with extra steps.', True),
    ' The shared shell gives us the real reuse without it.'
])

callout('TWO SEARCHES, ONE SHELL',
        'The interaction layer is built once and both surfaces consume it. The pipelines are built twice, on '
        'purpose, because they differ on corpus, execution, matching, ranking, offline behaviour and '
        'criticality — and because OPDS supplies a backend for one of them and nothing at all for the other. '
        'Voice input is simply a second way to produce a query string, so it plugs into the catalogue '
        'pipeline with no architectural change, and stays scoped to the catalogue.')

# ========================================================== 05 WEEKLY PLAN
secnum('05')
sectitle('Week-by-week plan')

callout('DEFINITION OF DONE',
        'Sections A to F working end to end against mock data by the close of Week 3, with the state-matrix '
        'pass complete. Section G — integration against real wokay and flambeau services — runs from Week 3 '
        'and closes in Week 4, gated on those services landing. team1 is never blocked: everything external '
        'is consumed through a stub we own.')

body_para([
    'Each week below is given per person, with the contingency stated alongside the deliverable. ',
    ('The contingency column is the point of these tables', True),
    ' — every block of work has a named fallback that keeps its owner moving when an external answer does '
    'not arrive, so a late dependency costs reshaping rather than idle time.'
])

h3('Week 1 — 10 to 14 August · Foundations, catalogue, search, institutions')
label('Days 1–2 · Foundation Sprint, all five people, one branch')
bullet('Design tokens transcribed from the design specification')
bullet('ContentItem, Institution, Session, DataAdapter and the resolveAccess signature committed as types')
bullet('Mock fixtures — ~40 items across three file formats and three tiers, 8 institutions, awkward cases included. DERIVED FROM wokay\'s Week 1 OPDS samples wherever they have landed, rather than hand-written from a guess')
bullet('App shell, bottom tab navigation, four empty routed screens merged to main')
bullet('Component inventory agreed and the library skeleton stubbed')

label('Days 3–5')
data_table(
    ['Member', 'Delivers', 'If the dependency slips'],
    [
        ['Prayas', 'F1, F2, A4, A8 — ONE OPDS adapter over one root feed and its groups, normalised to one ContentItem. No URN extraction: isbn is a top-level field. Next-link pagination only. Fixtures',
         'wokay\'s Week 1 samples not landed: build to their PUBLISHED field names, which we already have, and reconcile on arrival. Chase Monday morning — this is the one call that changes the week'],
        ['Moktik', 'Search shell and B1 — query state, templated-link expansion, rendering, next-link paging, empty and error states. Matching, tokenisation and ranking are WOKAY\'S, server-side',
         'The search endpoint is a WEEK 4 deliverable: build the pipeline against fixtures strictly behind the interface. Fetch-and-filter is a stand-in and must never become the design'],
        ['Khushi', 'Design tokens Day 1, then her five components and the state gallery',
         'None. This block has no external dependency at all, which is precisely why it goes first'],
        ['Keshav', 'B9, C1 — institution pipeline and list over wokay\'s PAGED SERVER endpoint, not client-side search. C2 detail moves to Khushi',
         'None on the schema — it is published in full. authRouting becomes an idpHint pass-through, since sign-in is always SAML and there is no branch to write'],
        ['Akriti', 'D1 skeleton — resolveAccess(item, session, loan, availability) as an ACQUISITION-LINK INTERPRETER, stubbed resolver, session shape; F6 pending-intent store; the error-code copy map',
         'Q-D is answered, so the tier is not a risk. The live one is W-22 work type: derive it in one function and record that screens 04 and 05 run on derived data'],
    ],
    [1100, 4400, 4292]
)

label('Coordination duty this week')
bullet([('Day 1', True), ' — ANSWER wokay\'s four asks, three of which are due this week. Institution list '
        'fields and sort order; subject facets, where the answer is YES; confirmation that we render the '
        'acquisition link; and the browse-free-content entry. Their silence-default on subject facets is not '
        'to build them, which deletes screen 01\'s Browse-by-Subject row.'])
bullet([('Day 1', True), ' — send the THREE remaining wokay questions: work type, DOI, and the section 3.1 '
        'conflict. Their document answers the rest, and sending answered questions costs credibility on '
        'these three.'])
bullet([('Day 1', True), ' — chase wokay\'s OpenAPI file, three OPDS samples and mock endpoints. Highest '
        'leverage item of the week: it closes our largest internal risk.'])
bullet([('Day 1', True), ' — escalate to leadership, ONE LETTER, TWO ITEMS: ratification of the corrected '
        'access matrix (L-6) bundled with the anonymous-feed reversal (L-2). The matrix half is the more '
        'urgent, and both are now backed by wokay\'s contract rather than by our preference.'])
bullet([('Day 3', True), ' — agree the sign-in handoff with flambeau. Narrower now: always SAML, and the '
        'payload is signIn.idpHint.'])
bullet([('By Friday', True), ' — resolve My Library ownership (L-7 / Q-Z) with t4targaryen and wokay, or '
        'carry it as an open escalation with a name against it.'])
bullet([('Friday', True), ' — hand the dependency board to next week\'s lead.'])

callout('WEEK 1 EXIT CRITERIA',
        'Catalogue browsable and searchable on mock data. Institution list, detail and search working. '
        'All 21 components merged with their state gallery. Contracts sent in writing to all four audiences, '
        'RECONCILED against wokay\'s published schema rather than re-asking what it already answers. '
        'wokay\'s four Week 1 asks answered. The work-type and DOI asks sent. L-6 and L-2 escalated together.')

h3('Week 2 — 17 to 21 August · Selection, access states, filters, detail')
data_table(
    ['Member', 'Delivers', 'If the dependency slips'],
    [
        ['Prayas', 'A0 shelf tabs read from the root feed\'s navigation rows; A1, A3, A5, A7, A9 catalogue home; D10, F8 loan-state fetching and caching; F9 items:batch',
         'The shelf set is per-institution CONFIGURATION and wokay can re-cut it without a rebuild, so the tab bar is data not code — mandatory rather than prudent. FL-9 unanswered: fetch loan state per item behind an aggressive cache, with the interface already shaped for bulk so switching is free'],
        ['Moktik', 'B3, B4, B6, B7, B8 filters, sort, result count and combined-state persistence; B11 voice input',
         'B5 access-tier filter is UNBLOCKED — accessTier already exists and is a free filter, so pull it forward from Week 3. The dimension genuinely at risk is SUBJECT, which needs facets wokay build only if we asked in Week 1'],
        ['Khushi', 'E1–E4 item detail across work types; D3 tier badges; screens 04 and 05. E5: the work type is DERIVED until wokay add the field',
         'W-13 metadata incomplete: render what exists and leave gaps blank rather than blocking on a full field set'],
        ['Keshav', 'C3–C8 selection, persistence, offline cache; screens 02, 06, 10',
         'FL-5 SAML return mechanism unknown: persist pending intent to disk rather than memory — the correct choice anyway if authentication leaves the app. W-13 metadata incomplete: render what exists, leave gaps blank'],
        ['Akriti', 'D1, D2, D4, D6, D7, D9, D11, A2, A6, F3 — the resolver, the loan rule, Borrow orchestration and the optimistic flip; screens 03, 18',
         'FL-3 entitlement access late: mock the session shape. This is the single largest stub in the plan and the one most likely to need rework'],
    ],
    [1100, 4400, 4292]
)

label('Coordination duty this week')
bullet('Take delivery of wokay\'s real institution directory — the first real dependency team1 receives from anyone — and delete the institution mock.')
bullet('Confirm the Read handoff with t4targaryen, and that Elite read-only is enforced at the reader too. Note it is already enforced twice: canPersist on the link, and DOWNLOAD_NOT_PERMITTED on the server.')
bullet('Chase W-22 work type. By the end of this week it is either answered or it is an escalation, because screens 04 and 05 are otherwise built on derived data indefinitely.')
bullet('Chase FL-9 bulk loan state — now the highest performance risk in the register, and the one item with no viable degraded mode.')

callout('WEEK 2 EXIT CRITERIA',
        'BOTH states — anonymous and signed-in institutional — walkable end to end on mock data, with the '
        'correct single action on every item and the pending-intent round trip working. Elite showing one '
        'button and Subscription showing two, which is the correction most likely to be got wrong. Filters '
        'live on file format and access tier, both free; subject live if the facets were asked for in Week 1.')

h3('Week 3 — 24 to 28 August · Hardening, state matrix, first integration')
data_table(
    ['Member', 'Delivers', 'If the dependency slips'],
    [
        ['Prayas', 'G1 — begin the real adapter swap; catalogue and pagination hardening',
         'Endpoints late: do fixture-parity work instead. Reshape fixtures to whatever wokay have described, so the swap is shorter when it comes'],
        ['Moktik', 'Search hardening and query-state depth. NOT a ranking quality pass — ranking is wokay\'s, server-side. B5 pulled forward to Week 2 because accessTier is a free filter',
         'No result count on OPDS feeds: count locally over the fetched set and display "40+" rather than a false total. And the real endpoint does not arrive until Week 4, so everything here is fixture-backed by design'],
        ['Khushi', 'F5 states everywhere; D8, D12, B10 rendered states; screens 14, 17; the state-matrix pass — her largest week',
         'None. It is scheduled here precisely because it is the one block that cannot slip on somebody else\'s account'],
        ['Keshav', 'Item detail hardening; offline institution cache; authentication round-trip hardening',
         'flambeau mock late: exercise the full round trip against our own placeholder, which we already built in Week 1'],
        ['Akriti', 'D13 Borrow failure handling and D14 the error-code copy map — replacing D5 Subscribe, which is deleted with the individual-subscriber scope; supports Khushi through the state-matrix pass',
         'FL-10 Borrow failure responses late: build the copy map against wokay\'s PUBLISHED code enumeration, which we already have, and confirm with flambeau that the codes reach us unchanged'],
    ],
    [1100, 4400, 4292]
)

callout('WEEK 3 EXIT CRITERIA — THE ONE THAT MATTERS',
        'team1 is feature-complete. Everything in Sections A to F is done, hardened and demonstrable on mock '
        'data. From this point on, anything outstanding is somebody else\'s service, not our feature work.')

h3('Week 4 — 31 August to 4 September · Integration close-out')
data_table(
    ['Member', 'Delivers', 'If the dependency slips'],
    [
        ['Prayas', 'G1 close-out; shape-mismatch fixes against the real feed', 'Demo on mocks and record the delta'],
        ['Moktik', 'G4 — BDD test suite', 'Unaffected; runs against mocks either way'],
        ['Khushi', 'G4 — accessibility pass', 'Unaffected'],
        ['Keshav', 'G2, G3 — real authentication routing and reader launch', 'Keep the placeholder round trip; it is demonstrable'],
        ['Akriti', 'G2 — real session and entitlements', 'Keep the mocked session; state clearly in the demo that it is mocked'],
    ],
    [1100, 4400, 4292]
)
body_para([
    'This week is ', ('deliberately underfilled', True), '. If everything landed on time it absorbs G4 early '
    'and we finish ahead. If nothing landed, this is the week we demonstrate on mocks and say so plainly '
    'rather than discovering the gap on the day.'
])

callout('THE ONE REAL RISK',
        'The three-week target holds only because we build entirely against mocks and consume every external '
        'service through a stub we own. It is threatened by exactly two things: mock shapes arriving late, '
        'and wokay being unable to serve an unauthenticated full catalogue. Both are answered by asking on '
        'Day 1, and both are already stubbed so that a late answer costs a day of reshaping, not a week of '
        'waiting.')

# ==================================================== 06 WORK DISTRIBUTION
secnum('06')
sectitle('Work distribution across the five')

body_para([
    ('Allocated by layer, not by feature.', True),
    ' Feature-vertical allocation is the intuitive choice and it is what produces the duplicated-component '
    'problem: five people each need a card, so five cards get written. Layer allocation gives each '
    'component and each contract exactly one author, and it keeps people out of each other\'s files.'
])

body_para([
    ('The team lead rotates weekly, so no lead role appears in this table.', True),
    ' Coordination is a duty that travels with the rota rather than a permanent assignment, and it carries '
    'its own reserved capacity — see below. All five carry a comparable build load. The OPDS audit in '
    'Section 03 is what set these boundaries — though note that its conclusion has since been REVERSED: '
    'catalogue search is server-side and entitlement-scoped, so matching, tokenisation and ranking are '
    'wokay\'s, not ours. Search remains a layer in its own right, but it is a query-surface layer rather '
    'than a search engine.'
])

h3('The five features')
data_table(
    ['Member', 'Feature owned', 'What that includes', 'Screens'],
    [
        ['Prayas Yadav', 'Catalogue, content & loan data',
         'A0 shelf tabs from the root feed\'s navigation rows, A1, A3, A4, A5, A7, A8, A9 · F1 OPDS normalisation to one ContentItem over ONE root feed and its groups, F2 adapters and fixtures, F9 items:batch, next-link pagination · D10 and F8 loan-state fetching, caching and invalidation · G1 the real-adapter swap',
         '01'],
        ['Moktik', 'Search & discovery',
         'B1–B13 · the query shell over wokay\'s server-side endpoint, every filter dimension as query parameters, year bucketing, sort, result count, combined-state persistence, voice input. Matching, tokenisation and ranking are WOKAY\'S — the ISBN exact-match path collapses to a top-level field read, and DOI has no backend field at all',
         '09, 11, 12'],
        ['Keshav Sharma', 'Institutions & sign-in — CAP-2, CAP-3',
         'C1–C9 · the institution query surface over wokay\'s PAGED SERVER endpoint, selection, persistence, change and clear, offline cache, and routing into flambeau by passing signIn.idpHint through — no auth-type branch, since sign-in is always SAML',
         '02, 06, 10'],
        ['Akriti Khetan', 'Access policy & borrowing',
         'D1, D2, D4, D6–D9, D11–D14 · resolveAccess as an ACQUISITION-LINK INTERPRETER, the corrected loan rule, Borrow orchestration and the optimistic flip, failure handling and the error-code copy map · F3 session context, F6 pending intent, F10 the typed error layer · A2, A6 feed scoping and the Access Gate. D5 Subscribe is deleted',
         '03, 18'],
        ['Khushi S Shukla', 'Item detail & every rendered state',
         'E1–E4 item detail across work types over Prayas\'s model, with the work type derived until wokay add the field (E5) · F4 app shell and navigation · F5 loading, empty, error and offline · D3 tier badges, D8 not_entitled and D12 waitlist — she renders every access state, Akriti decides when each applies · B10 search empty states · the Week 3 state-matrix pass',
         '04, 05, 14, 15, 16, 17'],
    ],
    [1500, 1900, 4700, 1692]
)

callout('THE ONE SEAM INSIDE FEATURE OWNERSHIP',
        'Access states are split between two people on purpose. Akriti decides WHEN a state applies — the '
        'resolver, the action matrix, the loan rule. Khushi renders WHAT it looks like — the tier badge, '
        'Access Restricted, the waitlist, and every loading, empty, error and offline surface. This is the '
        'only deliberate break in feature ownership, and it exists because the access feature was carrying '
        'three times the load of the lightest one. It works because the seam is the actions array: Akriti '
        'produces it, Khushi renders it, and neither needs to understand the other half.')

h3('Why feature ownership, and what it costs')
body_para([
    'An earlier version of this plan allocated by horizontal layer. ',
    ('This version allocates by feature, owned end to end', True),
    ' — each person owns their data access, their logic, their screens and their states. The reason is '
    'coordination cost: layer ownership means every feature is assembled from four people\'s work, and '
    'every question becomes a hand-off. Feature ownership means one person can finish something.'
])
body_para([
    ('The cost is real and worth naming.', True),
    ' Feature-vertical allocation is the classic cause of duplicated components and visual drift — five '
    'people each need a card, so five cards get written. Layer ownership prevented that structurally. '
    'This plan prevents it by timing instead: the component library is built by all five during the '
    'foundation phase and merged before any feature starts, so a feature owner always has something to '
    'compose from. That only holds if Phase 1 actually completes before Phase 2 begins. ',
    ('If the library slips into Week 2, the duplication problem returns immediately', True),
    ', and it will not be visible until Week 3.'
])
body_para([
    'Two things stay single-authored despite feature ownership: ', ('design tokens', True),
    ', written Day 1 before any component exists; and ', ('the shared type and adapter interfaces', True),
    '. If those fragment across five features, a contract change from wokay or flambeau becomes a '
    'five-file edit rather than a one-file edit — which is exactly the failure the layer model existed '
    'to prevent, and the one piece of it worth keeping.'
])

h3('How the features were balanced')
body_para([
    'Three facts drove it. ', ('The tabs decision', True),
    ' removed cross-feed dedup and cross-feed merge-and-rerank, which were the two heaviest blocks in the '
    'plan, bringing catalogue and search down to comparable size. ', ('The Borrow model', True),
    ' then added five deliverables and a mutable per-item state to the access feature, which is why loan '
    'work sits there rather than being spread. And ', ('experience is not level across the five', True),
    ': item detail and platform states go to the member still building React Native depth, because both '
    'are composition over other people\'s contracts, both are checkable against a finished design, and '
    'neither can be blocked by wokay or flambeau.'
])
body_para([
    'The lightest feature is institutions. It is no longer a client-side corpus — wokay serve it as a paged '
    'server endpoint — but it is still the smallest surface, and it has the earliest real dependency of '
    'anything on the team, in Week 2. It carries fewer screens as a result, and Keshav is the first person '
    'to absorb overflow if another feature runs long.'
])

h3('What the access layer actually owns, and when it shrinks')
body_para([
    'It is reasonable to ask whether this layer is redundant, given flambeau own authentication and wokay '
    'own the catalogue. It is not, but the reason is narrow and worth stating precisely. ',
    ('wokay know what an item is. flambeau know who the user is. Neither knows what this user may do with '
     'this item', True),
    ' — and that join can only happen where both facts meet, which is here. Four things live nowhere else: '
    'the join itself; the action policy (three tiers by three session states resolving to exactly one '
    'action, which is a product rule neither service holds); the pending-intent store, since neither '
    'service knows what the user tapped before signing in; and the seam that keeps flambeau\'s session '
    'shape and wokay\'s access field from leaking into every screen.'
])
callout('AND WHEN IT SHRINKS — THE CONTINGENCY',
        'This is the smallest code footprint on the team, and its size is a direct function of two open '
        'questions, and BOTH HAVE NOW BEEN ANSWERED IN THE DIRECTION THAT SHRINKS IT. W-7: wokay DO scope the '
        'feed to entitlements, so feed scoping (A2, A7) largely disappears. W-4: accessTier already exists — '
        'but the larger change is that resolveAccess no longer maps tier to action at all, it interprets the '
        'acquisition link, which is a smaller and more mechanical piece of code than the action matrix was. '
        'What does NOT shrink is the join, the pending-intent store, and the error-code copy map. Original '
        'note follows: resolveAccess '
        'degrades to close to a passthrough. Should both land that way — likely known by Wednesday of Week '
        '1 — this layer is under-loaded, and the holder absorbs, in order: E2 and E3 (article and '
        'journal/audio detail) from Keshav in Week 2, then the Week 3 state-matrix pass alongside Khushi. '
        'Both are named here so the reallocation is a decision already taken rather than one argued about '
        'mid-sprint.')

h3('Coordination is a rotating duty, not a role')
body_para([
    'The lead changes every week, which has one strong implication for this plan. ',
    ('The dependency board is the handover artefact.', True),
    ' Without it, each incoming lead restarts the chase from nothing, and the four blocking dependencies in '
    'Section 07 get asked three times and answered none. With it, handover is a five-minute read. Whoever '
    'holds the lead that week carries the board, the cross-team chasing and the escalations, and has a '
    'quarter of their week reserved for it. That reservation moves with the rota, not with the person, and '
    'their layer work for that week is planned accordingly.'
])

callout('THE FRIDAY HANDOVER',
        'Every Friday the outgoing lead brings the dependency board to a state where each item sits in '
        'exactly one of Asked, Answered, Stubbed or Integrated, with a date and a name against it. The '
        'incoming lead reads it and owns it. A weekly-rotating lead cannot survive without this ritual — it '
        'is the difference between five people sharing one thread of accountability and five people each '
        'holding a quarter of it for a week and dropping it.')

h3('Load by week')
body_para('Per-person deliverables and their contingencies are in Section 05, which is the operational view. '
          'This is the shape of it.')
data_table(
    ['Member', 'Week 1', 'Week 2', 'Week 3', 'Week 4'],
    [
        ['Prayas', 'His five components; F1, F2, A4 — ONE adapter, normalisation, fixtures, next-link pagination',
         'A0, A1, A3, A5, A7, A9 shelf tabs and catalogue home; D10, F8 loan-state fetching and caching; F9 items:batch',
         'G1 adapter swap; loan cache, offline behaviour and catalogue hardening',
         'G1 close-out; shape-mismatch fixes'],
        ['Moktik', 'His three components; search shell; B1 query surface — matching, tokenisation and ranking are wokay\'s',
         'B2–B4, B6–B8, B12 — exact match, filters, sort, count, state persistence; B11 voice',
         'Search hardening and query-state depth; supplies empty-state variants to Khushi. B5 pulled forward to Week 2, unblocked',
         'G4 — BDD suite'],
        ['Khushi', 'Tokens Day 1; her five components; F4 shell and navigation (paired); state gallery; C2 institution detail',
         'E1–E4 item detail, screens 04, 05; D3 tier badges',
         'F5 states everywhere; D8, D12, B10 rendered states; screens 14, 17; state-matrix pass',
         'G4 — accessibility pass'],
        ['Keshav', 'His four components; B9, C1 — institution list and its own query pipeline over wokay\'s paged endpoint. C2 detail moves to Khushi',
         'C3–C8 selection, persistence, offline cache; E1–E4 item detail, screens 04, 05',
         'Item detail hardening; auth round-trip hardening',
         'G2, G3 — real auth and reader launch'],
        ['Akriti', 'Her three components; D1 skeleton — signature, stub resolver, session shape; F6 pending intent',
         'D1, D2, D4, D6–D9, D11, A2, A6, F3 — the link interpreter, corrected loan rule, Borrow orchestration; screens 03, 18',
         'D13 Borrow failure handling; D14 the error-code copy map; supports Khushi on the state-matrix pass',
         'G2 — real session and entitlements'],
    ],
    [1100, 2223, 2223, 2223, 2023]
)

h3('File ownership, to keep merges cheap')
data_table(
    ['Area', 'Owner', 'Rule'],
    [
        ['theme / tokens', 'Khushi', 'Single author, written Day 1 before any component exists. Never split, whatever else is.'],
        ['components/*', 'Split five ways', 'Each component has one author (see Section 01). Anyone may add to the library; the original author reviews. A feature may never hold its own component.'],
        ['model/*, adapters/*', 'Prayas', 'Three adapters, one normalised model. The interface stays single-authored even under feature ownership; changes need Akriti as second reviewer.'],
        ['screens/detail/*', 'Khushi', 'Renders Prayas\'s model with Akriti\'s action bar. Must not reach past the adapter or compute access itself.'],
        ['access/*, session/*', 'Akriti', 'The only place licence, entitlement or action-selection logic may exist. Renders nothing — it emits an actions array.'],
        ['search/shell/*', 'Moktik', 'Consumed by both pipelines; Keshav is second reviewer.'],
        ['search/catalogue/*, filters/*, sort/*', 'Moktik', 'Independent.'],
        ['institutions/*', 'Keshav', 'Includes its own search pipeline. No OPDS backend exists for it.'],
        ['navigation/*', 'Khushi', 'F4. The classic conflict hotspot, so a single owner. Any change needs a second reviewer.'],
        ['loans/*', 'Prayas', 'Fetching, caching and invalidating loan state. The Borrow-versus-Read decision it feeds is Akriti\'s, in access/*.'],
        ['fixtures/*', 'Prayas', 'Anyone may add a case; nobody may reshape without Prayas.'],
        ['dependency board', 'That week\'s lead', 'Handed over every Friday. Never owned by a person for longer than a week.'],
    ],
    [2000, 1300, 6492]
)

# ================================================== 07 DEPENDENCY MAPPING
secnum('07')
sectitle('Dependency mapping')

h3('On wokay — CAP-1 Onboarding & Admin, CAP-5 OPDS Catalogue')
callout('MOST OF THIS TABLE IS NOW ANSWERED. DO NOT SEND IT AS IT WAS',
        'wokay have published a source of truth: seven collections, the catalogueItems record field by '
        'field, both app-facing REST shapes, the OPDS acquisition link in full, the error envelope, and a '
        'section 04 written specifically as the cross-team contract. Fourteen of the items below are '
        'answered by that document. Sending them anyway spends credibility on the three that are not. Read '
        'their section 04 before sending anything.')

body_para([
    ('And the direction of dependency has partly reversed.', True),
    ' wokay need four things FROM team1, three of them by end of Week 1, and each has a stated default if '
    'we are silent. See the second table. This inverts the plan\'s "silence is safe by design" premise: on '
    'those four, silence means THEY choose, and one of their defaults deletes a designed feature.'
])

data_table(
    ['#', 'We need', 'By', 'Blocks', 'Status'],
    [
        ['W-1', 'Unauthenticated catalogue feed', 'Wk 4', 'A1, whole anonymous model', 'ANSWERED, AND NARROWER. GET /opds/v1/public/catalogue exists with no token, but it is OPEN ACCESS ONLY — not the full catalogue. The full catalogue is entitlement-scoped and token-protected. Escalate to leadership as fact, not as a request'],
        ['W-2', 'Catalogue feed sample (OPDS 2.0 JSON)', 'Wk 1', 'F1, A1, B1', 'COMMITTED — OpenAPI file, THREE OPDS sample fixtures and mock endpoints, Week 1. This is the highest-leverage item in the whole register: it closes our own #1 risk. CHASE IT MONDAY MORNING'],
        ['W-3', 'Institution data schema and sample', 'Wk 2', 'C1, C2, B9', 'ANSWERED in full — id, code, name, type, country, city, logoUrl; plus branding, signIn.idpHint and catalogueUrl on detail'],
        ['W-4', 'A per-item field carrying access tier', 'Wk 1', 'B5, D1, D2, D3', 'ANSWERED BEFORE IT WAS ASKED. catalogueItems.accessTier is OPEN_ACCESS, SUBSCRIPTION or ELITE, exposed on items:batch, and offered as a FREE FILTER. The "hardest ask in the plan, longest lead time" required nothing'],
        ['W-5', 'Search contract', 'Wk 4', 'B1', 'ANSWERED — a templated search link, SERVER-SIDE, entitlement-scoped, and METADATA ONLY. Matching, tokenisation and ranking are theirs, not ours. Reverses the Section 03 audit'],
        ['W-6', 'Pagination contract', 'Wk 1', 'A4', 'ANSWERED — "follow the next link until it is absent. Do not count pages." OPDS is next-link only; page/size/total exists on the institutions endpoint alone'],
        ['W-7', 'Post-login feed behaviour', 'Wk 3', 'A2, A7, D1', 'ANSWERED — they scope it. The root feed is already filtered to what the institution is entitled to, plus open access, which appears for every institution regardless'],
        ['W-8', 'Stable item identifier', 'Wk 1', 'F6 pending intent', 'ANSWERED — prefixed strings, item_42'],
        ['W-9', 'DOI and ISBN URN prefixes', 'Wk 1', 'B2 exact-match', 'MOOT, AND REOPENED. isbn is a TOP-LEVEL field, so there is no URN to parse. But there is NO DOI ANYWHERE in their schema — see W-23'],
        ['W-10', 'Result count on every feed', 'Wk 2', 'B7', 'OPEN. total is returned by the institutions endpoint; unconfirmed on OPDS feeds. Count locally and approximate'],
        ['W-11', 'Sort mechanism — facet group or query parameter', 'Wk 2', 'B6', 'STILL OPEN. sort=publishedAt,desc is HINTED on the groups endpoint but not stated as the contract. Sort locally meanwhile'],
        ['W-12', 'Are subject, publisher and published consistently populated', 'Wk 2', 'B4 filters', 'OPEN. subjects and language are on the record; publishedYear has no clear equivalent — createdAt is ingest time, not publication'],
        ['W-13', 'Item metadata completeness', 'Wk 2', 'E1, E2', 'LARGELY ANSWERED by the published record. The two genuine gaps are work type and DOI'],
        ['W-14', 'Per-item declared formats', 'Wk 2', 'B3, D7', 'ANSWERED — assets[] carries format, mimeType and sizeBytes'],
        ['W-15', 'Institution sign-in-type field', 'Wk 2', 'C6', 'DEAD. "Institutional sign-in is ALWAYS SAML, so there is no authMethod field." We pass signIn.idpHint. authRouting becomes a pass-through'],
        ['W-16', 'Journal hierarchy representation', 'Wk 2', 'E3', 'MOOT — they model no work type at all, so there is no journal hierarchy to represent. Superseded by W-22'],
        ['W-17', 'Institution crest and logo URLs', 'Wk 2', 'C1', 'ANSWERED — logoUrl, plus branding.logoUrl and branding.primaryColor'],
        ['W-18', 'Endpoint-to-tab mapping for three feeds', 'Wk 1', 'A0, B3', 'THE PREMISE WAS WRONG. There are not three feeds. One root feed per institution carries navigation rows and groups; shelves are discovered AND configured per institution. Tab bar must be data, which we were already building'],
        ['W-19', 'Cross-feed duplicate identity', 'Wk 1', 'A0, F1', 'MOOT — one feed with groups, not three independent cursors. A work on two shelves is cheap, not a dedup problem'],
        ['W-20', 'Per-feed or global search link', 'Wk 1', 'B1, A0', 'ANSWERED — one search endpoint per institution, GET .../search?query=, scoped by entitlement rather than by shelf'],
        ['W-21', 'Where do articles live', 'Wk 1', 'A1, B3, E2, E3', 'MOOT AS ASKED — there is no Journals feed to put them in. Reopens as W-22'],
        ['W-22', 'WORK TYPE — a per-item field for book / journal / article / audiobook. Their contentType is PDF / EPUB / AUDIO, a FILE FORMAT, and they model no work type at all. Screens 04 and 05 differ by exactly this axis', 'Wk 1', 'E1, E2, E3, E5, A3', 'NEW, AND IT TAKES Q-D\'s DAY 1 SLOT. Same shape of ask as the access-tier field was, and it now has the longest lead time of anything we need. Derive it in one function meanwhile and say on the record that screens 04 and 05 run on derived data'],
        ['W-23', 'DOI. There is a top-level isbn field and no DOI anywhere in the schema. Article detail wants one', 'Wk 2', 'E2, B2', 'NEW. If it cannot be added, DOI comes off screen 04 — which should be a decision, not a blank field in the demo'],
        ['W-24', 'OPDS FACETS FOR SUBJECTS. Subjects are dynamic, so we cannot know which exist for an institution. Roughly half a day of theirs, and they build it ONLY IF WE ASK', 'Wk 1', 'A5 Browse-by-Subject, B4', 'NEW, AND WE MUST ANSWER YES IN WEEK 1. Their stated default if we are silent is to ship content-type and tier filters only, which deletes screen 01\'s Browse-by-Subject row and the SubjectChip component'],
    ],
    [500, 3200, 600, 2100, 3392]
)

h3('What wokay need from team1 — three of these are due in Week 1')
body_para([
    ('This table did not exist in the previous version, and it is the one with dates attached to our own '
     'name.', True),
    ' Each row carries wokay\'s stated default if we do not answer. Silence is not neutral here.'
])
data_table(
    ['They need', 'By', 'What they do if we are silent', 'Our answer'],
    [
        ['Which fields the institution list screen actually needs, and how we want it sorted', 'Wk 1',
         'Ship a superset and refine',
         'Their current shape is exactly right — id, code, name, type, country, city, logoUrl. Sort alphabetically by name ascending. Recently-used pinning is ours and client-side'],
        ['Whether we want subject filters. Yes means they add OPDS facets', 'Wk 1',
         'Ship content-type and tier filters only — WHICH DELETES A DESIGNED FEATURE',
         'YES. Screen 01 has a Browse-by-Subject row and we have a SubjectChip component; neither is buildable without facets'],
        ['Confirmation that we render the acquisition link rather than deciding buttons ourselves', 'Wk 1',
         'Document it as a contract test on their side',
         'Confirmed. resolveAccess interprets the link; accessTier drives the badge only. See D1 and Section 02'],
        ['Whether the home screen offers a browse-free-content entry beside find-your-institution', 'Wk 2',
         'They ship the public feed regardless; it costs us a button',
         'YES. With individual subscribers cut, the public feed IS our skip-institution path and it replaces the personal-account option on screen 03'],
    ],
    [2600, 500, 2700, 4092]
)

h3('On flambeau — CAP-6 Authentication')
data_table(
    ['#', 'We need', 'By', 'Blocks', 'If it is late'],
    [
        ['FL-1', 'Sign-in handoff contract — what we pass, how control returns', 'Wk 1', 'C6', 'Route to our own placeholder'],
        ['FL-2', 'MOOT — individual subscribers are cut from scope (wokay gate decision 7), so there is no B2B/B2C distinction to make. Every session is institutional', 'Wk 1', 'D5 removed, A2', 'None needed  — CLOSED'],
        ['FL-3', 'Session and entitlement access — how we read what the user is entitled to', 'Wk 2', 'A2, D1, F3', 'Mock the session shape'],
        ['FL-4', 'REMOVED — there is no Subscribe entry point, because individual subscribers are out of scope', '—', '—', '—'],
        ['FL-5', 'SAML return mechanism — does it leave the app, and how do we regain control. NARROWER NOW: SAML is the only path, so this is the only sign-in question that matters', 'Wk 1', 'F6 pending intent', 'Assume in-app webview with a deep-link return; persist pending intent to disk either way'],
        ['FL-6', 'MOOT — no individual users exist to bypass institution selection. The anonymous public feed is the skip path, and it needs nothing from flambeau', 'Wk 1', 'C7', 'None needed  — CLOSED'],
        ['FL-7', 'Session expiry behaviour while browsing', 'Wk 2', 'F3', 'Handle generically'],
        ['FL-8', 'Mock sign-in screen we can route into', 'Wk 2', 'C6', 'Placeholder'],
        ['FL-9', 'A BULK loan-state lookup — "which of these 40 item ids does this user currently have on loan?" One call per page of results, not one per item', 'Wk 1', 'D10, F8, every catalogue surface', 'Fetch per item and cache aggressively. This degrades badly and is the likeliest source of a slow feed in the demo  — NEW, and the highest-risk of these'],
        ['FL-10', 'The Borrow endpoint itself — params, and the full set of failure responses: seat limit reached, licence expired, already on loan, not entitled', 'Wk 2', 'D9, D13', 'Stub to the agreed contract; render a generic failure  — NEW'],
        ['FL-11', 'GET /api/v1/availability?itemId= — ELITE ONLY, and the app calls it, never wokay. The feed carries copies.total and never copies.available, because wokay do not call flambeau while building a feed', 'Wk 4', 'D12', 'THE CONTINGENCY IS NOW THE DESIGN. no_seats is detail-screen-only and a catalogue card can never render it. If this endpoint slips, the feed still ships copies.total and the app shows no live count  — CONFIRMED AND NARROWED'],
        ['FL-12', 'Loan duration and expiry, and what we do when a loan lapses while the user is reading or offline', 'Wk 2', 'F8, D13', 'Treat an expired loan as no loan and re-show Borrow  — NEW'],
        ['FL-13', 'Confirmation of the corrected loan model: open access needs NO loan in any session state; Subscription writes a loan and skips Redis because it is unlimited; Elite acquires a Redis lease first, then writes the loan, or writes a HOLD and returns a queue position. wokay state this as flambeau\'s branch on accessLevel', 'Wk 1', 'D9, D11, D12', 'wokay\'s document already specifies it; proceed on that and ask flambeau only to confirm  — NEW, and it replaces the loan-rule confirmation that was wrong'],
        ['—', 'REMOVED: Buy, and the Subscribe entry point, with individual subscribers out of scope entirely. Borrow has returned with a different meaning — see D9 and Section 02.', '—', '—', '—'],
    ],
    [500, 3200, 600, 2100, 3392]
)

h3('On t4targaryen — CAP-7 Reader & Offline')
data_table(
    ['#', 'We need', 'By', 'Blocks', 'If it is late'],
    [
        ['T-1', 'Reader launch contract for the Read action', 'Wk 2', 'D6', 'Stub'],
        ['T-2', 'Download handoff contract, and who owns the download store', 'Wk 2', 'D7', 'Stub  — NEW'],
        ['T-3', 'Confirmation that Elite is read-only at the reader too, not just in our UI', 'Wk 2', 'D4', 'We simply never offer Download  — NEW'],
        ['T-4', 'Ownership of the downloaded/offline indication on catalogue items', 'Wk 2', 'A3', 'Omit initially'],
        ['T-5', 'Guest downloads — do they exist, and who stores them', 'Wk 2', 'D7', 'Omit  — NEW'],
    ],
    [500, 3200, 600, 2100, 3392]
)

h3('On leadership')
data_table(
    ['#', 'We need', 'By', 'Blocks', 'If it is late'],
    [
        ['L-1', 'Ratification that team1 owns catalogue browse and search UI', 'Wk 1', 'Sections A, B, D, E', 'Escalate — most of this plan depends on it. Note wokay\'s section 03 independently lists team1 as owning "find institution, browse, search, detail, my library", which supports it'],
        ['L-2', 'Ratification that the anonymous feed is OPEN ACCESS ONLY, which reverses Design Specification 4.1 and step 2 of the flow in index.html', 'Wk 1', 'A1, A2, whole model', 'CHARACTER CHANGED. This is no longer "will leadership ratify our reversal" — wokay CANNOT SERVE the v1 model. The full catalogue is entitlement-scoped and token-protected; the unauthenticated feed is open-access-only, full stop. Still escalate, but as a fact backed by the backend team\'s published contract. That is a stronger letter — and BUNDLE IT with L-6'],
        ['L-3', 'Confirmation that Buy AND Subscribe are out of the prototype, along with individual subscribers entirely (wokay gate decision 7). Borrow survives with a different meaning — loan creation on entitled content', 'Wk 1', 'D5, D9', 'LARGELY SETTLED by wokay\'s gate decision. Proceed with all three removed  — CHANGED'],
        ['L-4', 'One shared app, or one per team', 'Wk 1', 'F4', 'Build our own; keep it portable. Note that wokay state plainly that team1 has NO backend module — React Native only — which settles the backend half of this'],
        ['L-5', 'Ratification of shelves-as-tabs on screen 01, which still corrects the single mixed list in the mockup', 'Wk 1', 'A0, B1, F1', 'Build tabs from a config list. STAYS OPEN but narrows: there are not three fixed feeds to merge or split, there is one root feed whose shelf set is per-institution configuration, so the tab bar must be data regardless of how this is ratified'],
        ['L-6', 'RATIFICATION OF THE CORRECTED ACCESS MATRIX. Design Specification 3.1 is a signed document and its matrix is inverted on Elite and on Open Access. It has been corrected to v0.3 against wokay\'s section 02', 'Wk 1', 'D1, D2, D9, D12, every screen', 'NEW, AND MORE URGENT THAN L-2. Until this is ratified, two contradictory access matrices are in circulation and nobody can write the resolver against a single agreed source. Build to v0.3 and say so in writing'],
        ['L-7', 'MY LIBRARY OWNERSHIP. wokay\'s section 03 lists team1 as owning "my library" and pitches items:batch at that screen; our plan assigns screens 07 and 08 to t4targaryen as one CAP-7 block', 'Wk 1', 'F4 component library, screen 08, Khushi\'s block', 'NEW. Our position: screen 08 is the offline shelf and everything decision-bearing on it — download progress, bytes on disk, the keystore, guest-download persistence — is CAP-7. Resolve by Friday of Week 1 or carry it as an open escalation. Adopt items:batch into the data layer either way'],
    ],
    [500, 3200, 600, 2100, 3392]
)

callout('THE FOUR THAT DECIDE WHETHER THIS PLAN HOLDS — REVISED',
        'Three of the previous four are answered, and the list is now shorter and different. W-2, the Week 1 '
        'OPDS samples and OpenAPI file — this is the one that closes our largest internal risk, and it is '
        'the only item where chasing on Monday morning changes the week. W-22, the work type field — it '
        'takes W-4\'s old slot as the longest-lead ask, because screens 04 and 05 differ by an axis wokay do '
        'not model. W-24, subject facets — not because it is hard, but because their default if we say '
        'nothing is not to build it, which deletes a designed feature. L-6, ratification of the corrected '
        'access matrix — because until it lands, two contradictory matrices are in circulation. '
        'W-4 required nothing. W-1 is answered and narrower. L-2 is now a fact rather than a request.')

# =================================================== 08 AMBIGUITY REGISTER
secnum('08')
sectitle('Ambiguity register, and how the plan absorbs it')

body_para([
    'Every open question below has an owner, a stub already shipping, and a date after which we stop waiting '
    'and escalate. ',
    ('The purpose of the stub is that no ambiguity ever appears on the critical path.', True),
    ' A late answer costs us a day of reshaping, not a week of blocked work — and that is the only reason a '
    'three-week target is credible with this many unknowns.'
])

h3('The register')
body_para([('Owner here means owner of the answer, not owner of the escalation.', True),
           ' Escalation belongs to whoever holds the lead that week.'])
data_table(
    ['#', 'Open question', 'Affects', 'Owner', 'Stub we ship now', 'Escalate by'],
    [
        ['Q-A', 'CLOSED — a user reaches a non-entitled item through pending intent. It renders not_entitled with no action, screen 14. Simplified further now that Subscribe is deleted: there is no alternative resolution', 'D8, F6', '—', 'Shipped as answered', 'Closed'],
        ['Q-B', 'CLOSED, AND NARROWER — wokay serve GET /opds/v1/public/catalogue with no token, but it is OPEN ACCESS ONLY. There is no unauthenticated full catalogue', 'A1, whole anonymous model', '—', 'Anonymous surface built against the public feed only', 'Closed'],
        ['Q-C', 'CLOSED — wokay scope the feed. The root feed is already filtered to the institution\'s entitlements, plus open access', 'A2, A7, D1', '—', 'Shipped as answered', 'Closed'],
        ['Q-D', 'CLOSED — catalogueItems.accessTier already exists, is exposed on items:batch, and is offered as a free filter. The highest-stakes item in the register required nothing', 'B5, D1, D3, the action matrix', '—', 'Shipped as answered', 'Closed'],
        ['Q-E', 'CLOSED — a templated search link, server-side and entitlement-scoped. Fetch-and-filter survives only as a fixture stand-in behind the pipeline interface, and must never become the design', 'B1', '—', 'Fetch-and-filter behind the interface, Week 1 only', 'Closed'],
        ['Q-F', 'Is a result count available on OPDS feeds? total IS returned by the institutions endpoint', 'B7', 'Moktik', 'Count locally and approximate on catalogue surfaces', 'Fri, Wk 2'],
        ['Q-G', 'CLOSED, AND REOPENED AS Q-X — isbn is a top-level field, so no URN parsing. But there is no DOI anywhere in wokay\'s schema', 'B2 exact-match', '—', 'Read isbn directly; DOI omitted', 'Closed'],
        ['Q-H', 'Sort — a "Sort by" facet group, or a query parameter? sort=publishedAt,desc is hinted on the groups endpoint but not stated as the contract', 'B6', 'Moktik', 'Sort locally over the fetched page', 'Fri, Wk 2'],
        ['Q-I', 'CLOSED — individual subscribers are cut from scope (wokay gate decision 7). There are no B2C users to be entitled to anything', 'A2, D5', '—', 'Institutional sessions only', 'Closed'],
        ['Q-J', 'CLOSED, AND NARROWED — no_seats is live but ELITE ONLY, because Elite is the only tier with a finite copy count. Subscription is unlimited, so no queue can ever form on it', 'D12, screen 14', '—', 'Shipped as answered', 'Closed'],
        ['Q-K', 'Who owns Profile (screen 10) in full?', 'F7', 'Keshav', 'Build the institution section as a standalone route', 'Fri, Wk 1'],
        ['Q-L', 'CLOSED AS ASKED — wokay model no work type at all, so there is no journal hierarchy. Reopens as Q-Y', 'E3, F1', '—', 'Flat', 'Closed'],
        ['Q-M', 'Guest downloads — do they persist, and whose screen?', 'D7', 'Keshav', 'Hand to t4targaryen, no local store. Related to Q-Z', 'Fri, Wk 2'],
        ['Q-N', 'Does selecting an institution while anonymous scope anything before sign-in? Note the institution endpoints need NO token, so selection is possible while anonymous', 'C3', 'Keshav', 'No — selection takes effect only after auth', 'Wed, Wk 2'],
        ['Q-O', 'One shared app across the four teams, or one each?', 'F4', 'Akriti', 'Our own app, kept portable. wokay confirm team1 has NO backend module, which settles half of it', 'Fri, Wk 1'],
        ['Q-P', 'Does SAML leave the app, and how does control return? Now the ONLY sign-in question, since SAML is the only path', 'F6, C6', 'Keshav', 'Persist pending intent to disk; in-app webview with deep-link return', 'Thu, Wk 1'],
        ['Q-Q', 'THE PREMISE WAS WRONG — there are not three content-type feeds. One root feed per institution carries navigation rows and groups, and the shelf set is per-institution configuration', 'A0, B3', '—', 'Tab bar driven from the root feed\'s navigation rows — data, never a hardcoded list', 'Closed'],
        ['Q-R', 'CLOSED — one feed with groups rather than three cursors, so a work on two shelves is cheap and there is no cross-feed dedup problem', 'A0, F1', '—', 'Render duplicates if they occur', 'Closed'],
        ['Q-S', 'Has leadership ratified shelves-as-tabs? It still corrects the single mixed list on design screen 01', 'A0', 'Akriti', 'Build tabs from a config list — now mandatory, since the shelf set is the institution\'s to configure', 'Fri, Wk 1'],
        ['Q-T', 'CLOSED AS ASKED — there is no Journals feed for articles to sit inside. Reopens as Q-Y', 'A1, B3, E2, E3', '—', 'n/a', 'Closed'],
        ['Q-U', 'Does flambeau expose a BULK loan-state lookup, or only per-item? A page of forty results cannot make forty calls. Note wokay ship items:batch for the metadata half of this problem, which is a precedent to cite', 'D10, F8, every catalogue surface', 'Prayas', 'Per-item fetch behind an aggressive cache, with the interface shaped for bulk so switching is free', 'Tue, Wk 1  — still the highest performance risk'],
        ['Q-V', 'CLOSED — the oddity was an artefact of an incorrect loan rule. Open access resolves identically in every session state: Read and Download, no loan, no borrow step', 'D9, the action matrix', '—', 'Shipped as answered', 'Closed'],
        ['Q-W', 'What shows when loan state cannot be reached, offline or flambeau down — Borrow, Read, or a disabled button?', 'F8, D10', 'Prayas', 'Last cached state; if never cached, show Borrow and fail on tap', 'Fri, Wk 2'],
        ['Q-X', 'DOI. There is a top-level isbn field and no DOI anywhere in wokay\'s schema. Article detail wants one', 'E2, B2', 'Prayas', 'Omit DOI from screen 04 until answered — a decision, not a blank field', 'Wed, Wk 1  — NEW'],
        ['Q-Y', 'WORK TYPE. wokay model none — their contentType is PDF / EPUB / AUDIO, a file format. Screens 04 and 05 differ by exactly the axis they do not have', 'E1, E2, E3, E5, A3', 'Akriti', 'Derive it in ONE function; both screens run on derived data and we say so on the record', 'Tue, Wk 1  — NEW, takes Q-D\'s slot as the longest-lead ask'],
        ['Q-Z', 'MY LIBRARY OWNERSHIP. wokay\'s section 03 lists team1 as owning "my library" and pitches items:batch at that screen; our plan assigns screens 07 and 08 to t4targaryen as one CAP-7 block', 'F4 component library, screen 08, Khushi\'s block', 'Akriti', 'Hold at 21 components with screen 08 excluded; adopt items:batch into the data layer regardless, since it is the right shape either way', 'Fri, Wk 1  — NEW'],
        ['Q-AA', 'Do we theme from branding.primaryColor? wokay carry a per-institution primary colour', 'F4, every component', 'Khushi', 'DECIDED BY TEAM1: no. One fixed palette; the field is carried and unused, so adopting it later is additive rather than a refactor', 'Decided  — NEW'],
        ['Q-AB', 'Does the shelf set change per institution, and can wokay re-cut it mid-project? Their admin console has a per-institution catalogue config for feed title, shelf order and page size', 'A0, F1', 'Prayas', 'Tab bar is data. A re-cut costs nothing, which is the whole reason for building it that way', 'Fri, Wk 1  — NEW'],
    ],
    [500, 2900, 1300, 800, 2600, 1692]
)

callout('WHAT THE SOURCE OF TRUTH DID TO THIS REGISTER',
        'Twelve of the twenty-two open questions closed at once, and two closed by having their premise '
        'removed rather than answered. Five new ones opened. The net is a shorter register with a different '
        'centre of gravity: it was dominated by "will wokay expose X", and it is now dominated by "wokay '
        'model no work type", "who owns My Library", and one endpoint arriving in Week 4. That is a better '
        'register to be holding on day one of five, and it is the argument for reading another team\'s '
        'document before sending them questions.')

h3('Reserved capacity — the buffer that makes the register survivable')
body_para('This capacity is allocated in advance and is not counted as available for feature work. Removing '
          'it is the fastest way to turn a one-day answer into a one-week slip.')
data_table(
    ['Who', 'Reserved', 'Held against'],
    [
        ['That week\'s lead', '25% of their week, every week', 'Dependency chasing, contract negotiation, escalation and the Friday handover. This moves with the rota, not with a person — whoever holds it has their layer work for that week planned around it.'],
        ['Prayas', '20% in Weeks 2 and 3, DOWN from 30%', 'Shape mismatch on the real-adapter swap, plus loan-state caching. Six of the seven questions that used to land in his layer are closed — Q-C, Q-G, Q-L, Q-Q and Q-R all went, and the three-feed dedup and merged-pagination work went with them. What remains is Q-U, Q-W and Q-X, plus FL-9, which could still force a redesign of how loan state is fetched. If wokay\'s Week 1 samples land, reduce this again.'],
        ['Moktik', '25% in Weeks 2 and 3, UP from 20%', 'No longer the largest engine in the plan — matching, tokenisation and ranking are wokay\'s. But the reserve goes UP rather than down, because his dependency now arrives in WEEK 4, which is team1\'s integration and BDD week. He builds against fixtures for three weeks and integrates in the busiest one. Q-F and Q-H still change how a finished piece behaves; Q-D no longer blocks B5 at all.'],
        ['Keshav', '20% in Week 2', 'Two areas converging in one week — authentication churn (FL-1, FL-5, FL-8, Q-P) alongside item detail. Reduced somewhat by there being no auth-type branch to build: sign-in is always SAML and authRouting becomes a pass-through.'],
        ['Khushi', '25% in Week 3, plus review time throughout', 'The state-matrix pass across screens she did not build, and now every rendered access state as well. She is the member still building depth, so her work is deliberately the most reviewable and the least externally blocked — only three dependencies touch it.'],
        ['Akriti', '25% in Week 2, plus the coordination duty', 'Still the most externally blocked feature by a wide margin. FL-3 and FL-10 landing late means reworking the resolver rather than extending it — and the resolver has ALREADY been reshaped once, from a tier-based rules engine into an acquisition-link interpreter, which is the change this reserve exists for. Also holds Q-Y, Q-Z and the L-6 escalation.'],
        ['Whole team', 'Week 4 underfilled by design', 'Late external services. If everything lands on time, Week 4 absorbs G4 early and we finish ahead.'],
    ],
    [1300, 1700, 6792]
)

h3('Re-planning when the other teams move')
body_para([
    'Every fallback above is per-item: one answer arrives late, one stub absorbs it. That is not enough on '
    'its own, because the other three teams will not move at the pace this plan assumes — some things will '
    'land early, others will not land at all. ',
    ('What follows are the points at which the plan re-cuts itself, decided now rather than argued about '
     'mid-sprint.', True)
])

data_table(
    ['Checkpoint', 'What we assess', 'If ahead', 'If behind'],
    [
        ['Mon, Week 1', 'W-2 — have wokay\'s OpenAPI file, three OPDS samples and mock endpoints landed',
         'Derive the fixtures from their samples instead of hand-writing them, and close our largest internal risk on day one',
         'Hand-write fixtures against their published field names, which we already have, and reconcile on arrival. Chase daily — this is the single item where a Monday phone call changes the week'],
        ['Wed, Week 1', 'W-22, W-24, FL-9, L-6 — the work-type field, subject facets, bulk loan state, and ratification of the corrected access matrix',
         'Drop the corresponding stubs and build against the real shapes immediately; Week 3 integration starts a week early',
         'Every one has a stub already committed. No work stops. Escalate W-22 and FL-9 specifically — they are the two with no viable degraded mode. W-24 is not a degraded mode but a deleted feature, so answering it is ours to do rather than theirs'],
        ['Fri, Week 1', 'Have contracts arrived in writing from all four audiences',
         'Begin shaping fixtures to the real contracts over the weekend hand-over',
         'Freeze our assumed contracts and declare them the working spec. Silence becomes acceptance; say so explicitly in writing'],
        ['Fri, Week 2', 'Is a flambeau mock sign-in reachable, and does Borrow have an endpoint',
         'Wire the real round trip and free the Week 3 buffer for the state-matrix pass',
         'Ship the placeholder round trip. It is demonstrable, and it was built in Week 1 for exactly this'],
        ['Wed, Week 3', 'Are real endpoints live for catalogue, session and loans',
         'Pull G1 and G2 forward; use Week 4 for polish rather than catch-up',
         'Hold on mocks and freeze. Week 4 becomes integration, and the demo runs on mocks with that stated plainly rather than discovered on the day'],
        ['Any time', 'A capability turns out to be unowned — as Buy was, and as loan expiry may be',
         'Absorb it if it is small and ours by adjacency',
         'Stub it, name it in the register, escalate the same day. Do not silently absorb another team\'s scope — it is the fastest way to miss our own'],
    ],
    [1300, 2900, 2700, 2892]
)

callout('THE ASYMMETRY TO PLAN AROUND',
        'Another team being late costs us a stub we have already written. Another team being early costs us '
        'nothing. But another team changing a contract after we have integrated against it is the expensive '
        'case, and it is the one no stub protects. That is why fixtures are updated before code on every '
        'change, and why the adapter and session interfaces are owned by single named people — so a '
        'contract change is one file, not fifteen.')

h3('The rule that ties it together')
callout('STUB-FIRST',
        'Every external dependency is consumed through an interface we define, with a stub committed on Day '
        '1 and a fixture behind it. No team member is ever blocked — they are running on a stub. When a real '
        'answer arrives, the fixture is updated first, we observe what breaks, and only then do we integrate. '
        'This is what converts thirteen open questions from a schedule risk into a day of work each.')

# ======================================================= 09 QUESTIONS
secnum('09')
sectitle('Questions and deliverables from the other teams')

body_para([
    'Send on Day 1 of Week 1. Each list leads with the request that actually unblocks us — sample data and '
    'contracts, not finished services. ',
    ('Attach our assumed contract as a typed shape to each list', True),
    ' and ask them to correct it rather than answer from a blank page.'
])

h3('For wokay')
callout('REWRITTEN. THE PREVIOUS TWENTY-THREE QUESTIONS ARE OBSOLETE',
        'wokay published a source of truth and its section 04 is written as the cross-team contract. It '
        'answers fourteen of the twenty-three questions this section used to carry, and removes the premise '
        'of three more. Sending the old list would spend credibility on questions their own document '
        'answers, immediately before asking them for the three things it does not. What follows is: what we '
        'answer for them, then what genuinely remains.')

label('What we OWE wokay — three of these are due by end of Week 1')
body_para([
    ('This is the part that did not exist before, and it is the part with our name and a date on it.', True),
    ' Each has a stated default if we say nothing, and on subject filters their default deletes a feature '
    'the design already specifies. Silence is not neutral.'
])
bullet([('Institution list fields and sort order', True), ' — your current shape is exactly right: id, code, '
        'name, type, country, city, logoUrl. Please do not add to it. Sort alphabetically by name, '
        'ascending, as the server default. Recently-used pinning is ours and client-side, so it needs '
        'nothing from you. On detail we additionally use branding.logoUrl, signIn.idpHint and catalogueUrl, '
        'all of which you already return.'])
bullet([('Subject filters — YES, please build the OPDS facets', True), '. Screen 01 has a Browse-by-Subject '
        'row and we have a component for it. Subjects are dynamic, so we cannot know which exist for an '
        'institution from a page of twenty results. We understand this is roughly half a day of yours, and '
        'we would rather spend it than lose the feature.'])
bullet([('Acquisition-link rendering — confirmed', True), '. We render from the acquisition link and its '
        'properties. rel=open-access gives Read and Download; rel=acquisition with canPersist true gives '
        'Borrow then Read and Download; rel=borrow with canPersist false gives Borrow then Read only; no '
        'link means no button. resolveAccess does not read accessTier to decide an action — the tier drives '
        'the badge and nothing else. A contract test on your side is welcome anyway.'])
bullet([('Browse-free-content entry — yes', True), '. With individual subscribers cut, /opds/v1/public/'
        'catalogue IS our second entry point, and it replaces the personal-account option on our Access Gate '
        'sheet. It is not a nice-to-have.'])

label('What we need delivered')
bullet([('Week 1 — the OpenAPI file, three OPDS sample fixtures and the mock endpoints.', True),
        ' You say these matter more in Week 1 than any endpoint you could actually build, and you are right: '
        'they close the largest internal risk in our plan, which is our adapter author designing a '
        'normalisation model blind. This is the one item we will chase on Monday morning.'])
bullet([('Week 2', True), ' — the real institution directory, so we can stop using mocks. This is the '
        'earliest real dependency any team1 member gets from anyone.'])
bullet([('Week 3', True), ' — the first real entitlement-scoped root feed and publication detail.'])
bullet([('Week 4', True), ' — search, shelf paging, the public feed and items:batch.'])

label('Questions — there are three, and one of them is a conflict rather than a question')
for i, q in enumerate([
    'WORK TYPE, and this is now our longest-lead ask. Your contentType is PDF, EPUB or AUDIO, which is a '
    'file FORMAT, and we consume it as one. Separately, our screens 04 and 05 are article detail and book '
    'detail — they differ by work type: book, journal, article, audiobook. Your model does not carry that '
    'axis at all. Can you add it? This is the same shape of request the access-tier field would have been, '
    'except that one turned out to already exist. We have put the field in our fixtures and derive it '
    'heuristically meanwhile, isolated in one function — please correct our shape rather than starting from '
    'a blank page. If it cannot be added, tell us in Week 1 and we will stop distinguishing the two screens.',

    'DOI. You have a top-level isbn field, which is simpler than we assumed and removes a URN-parsing job '
    'from our side — thank you. But there is no DOI anywhere in the schema, and article detail wants one. Is '
    'identifiers.doi addable? If not, DOI comes off screen 04, which we would rather decide than discover.',

    'A CONFLICT WE CANNOT RESOLVE OURSELVES, flagged because it is your contract being treated as correct. '
    'Design Specification section 3.1 — a signed document — carries an access matrix that contradicts your '
    'section 02 on two of three tiers: it states that Open Access requires a Borrow once signed in, and '
    'that Elite never borrows. Your document states the opposite on both, and we believe your document. We '
    'have corrected the design specification to v0.3 and escalated to leadership for ratification. We are '
    'telling you because until it is ratified there are two contradictory access matrices in circulation, '
    'and nobody on this project can write an access resolver against a single agreed source.',
], start=1):
    p = para(space_after=4)
    run(p, f'{i}.  ', bold=True, color=NAVY)
    p.add_run(q)

label('Confirmations only — no answer needed unless we have read you wrong')
body_para([
    'Everything below we have taken from your document and built against. ',
    ('Correct us only where we are wrong; silence here we will read as agreement.', True)
])
for i, q in enumerate([
    'There is ONE root catalogue feed per institution carrying navigation rows and groups, not three feeds '
    'split by content type. Shelves are discovered from the root feed and configured per institution through '
    'your catalogue config, so our tab bar is driven from data and a re-cut on your side costs us nothing.',
    'We follow hrefs and never build URLs. We take catalogueUrl from institution detail and every later hop '
    'from a feed. The only construction we do anywhere is expanding ".../search{?query}".',
    'We paginate by following the next link until it is absent and we do not count pages. We understand '
    'page, size and total exist only on the institutions endpoint.',
    'A zero-result search returns a feed with a navigation entry rather than an empty publications array, '
    'and we render that as a browse-instead affordance rather than treating a missing key as an error.',
    'Catalogue search is server-side, entitlement-scoped and searches metadata only — title, authors, '
    'subjects, description — not the text inside books. We are putting that in our UI copy so a reviewer '
    'typing a phrase from page 88 does not file it as a bug.',
    'Filters are query parameters applied before pagination, never client-side within a fetched page. '
    'contentType and accessTier are fixed enums and free; we hardcode the chips and send the parameter.',
    'accessTier on the item is OPEN_ACCESS, SUBSCRIPTION or ELITE, uppercase, and we use it for the badge '
    'only.',
    'The feed carries copies.total and never copies.available, so we fetch availability from flambeau '
    'per-item on the detail screen and a catalogue card never shows a live count or a waitlist state.',
    'Institutional sign-in is always SAML, there is no authMethod field, and we pass signIn.idpHint '
    'through to flambeau untouched.',
    'Item ids are stable prefixed strings such as item_42, which is what our pending-intent store holds '
    'across the sign-in round trip.',
    'items:batch is capped at 100 ids and returns items, notFound and denied as three separate keys. We '
    'treat denied and notFound differently rather than collapsing both into an error.',
    'Errors arrive in one envelope — timestamp, status, code, message, path, traceId — and we key our error '
    'copy on the CODE rather than the HTTP status, so a denial can say the library subscription has expired '
    'rather than not available. We also understand NOT_FOUND is deliberately indistinguishable between '
    'unknown, archived and not-entitled, and our copy will not guess.',
    'An inactive institution returns 404 rather than 403 so its existence is not disclosed, and we treat it '
    'as not-found.',
    'team1 has no backend module. React Native only.',
    'ONE THING TO CORRECT ON YOUR SIDE: your section 03 lists team1 as owning "my library". Screens 07 and '
    '08 are the reader and the offline shelf, and everything decision-bearing on screen 08 — download '
    'progress, bytes on disk, the keystore, guest-download persistence — is t4targaryen\'s CAP-7. We are '
    'adopting items:batch into our data layer regardless, because it is the right shape for any surface '
    'holding item ids without metadata. But the screen is not ours, and we are raising it with '
    't4targaryen in the same week.',
], start=1):
    p = para(space_after=4)
    run(p, f'{i}.  ', bold=True, color=NAVY)
    p.add_run(q)

h3('For flambeau')
label('What we need delivered')
bullet([('Week 1', True), ' — the sign-in handoff contract: what we pass in, how control returns, where the '
        'user lands.'])
bullet([('Week 1', True), ' — confirmation of the corrected loan model, per wokay\'s accessLevel branch: '
        'OPEN_ACCESS needs no loan at all; ENTITLED_UNLIMITED writes a loan and skips Redis; '
        'ENTITLED_CONCURRENT acquires a lease first, then writes the loan, or writes a hold.'])
bullet([('Week 2', True), ' — a mock sign-in screen we can route into, and a way to read the signed-in '
        'user\'s session and entitlements.'])
bullet([('Week 4', True), ' — GET /api/v1/availability?itemId= for Elite copy counts, which the app calls '
        'and wokay never do. Without it the feed ships copies.total and we show no live count.'])
bullet([('REMOVED', True), ' — the Subscribe entry point. Individual subscribers are out of scope per '
        'wokay\'s gate decision 7, so there is no Subscribe flow to contract.'])

label('Questions')
body_para([
    ('Several of these are narrower than they were, because wokay\'s document settles the backend half.', True),
    ' Institutional sign-in is always SAML, so there is no auth-type branch to agree; individual '
    'subscribers are cut, so the B2B/B2C questions are gone.'
])
for i, q in enumerate([
    'We pass signIn.idpHint, taken from wokay\'s institution detail response, plus the institution '
    'identifier. Since sign-in is always SAML there is no auth type to route on. Is idpHint what you expect, '
    'and in what shape?',
    'Do you want one entry point that branches internally, or separate screens we route to? One entry point '
    'means we cannot get the routing wrong.',
    'How does control return to us after sign-in, and where does the user land — our catalogue, or a screen '
    'you own? We need to land them back on the exact item they tapped before signing in.',
    'Does SAML or SSO leave the application — system browser, in-app webview — and what is the return '
    'mechanism? This determines whether our pending-intent store survives the round trip, or whether we '
    'need to persist it to disk.',
    'How do we read the signed-in user\'s session to scope the catalogue? Note we do NOT need entitlement '
    'data from you — wokay scope the root feed themselves, so we need identity and institution, not a list '
    'of what they may read.',
    'Confirm the session payload shape. With individual subscribers cut, we have reduced it to { userId, '
    'institutionId, roles[], exp } and dropped the B2B/B2C type field, since every session is now '
    'institutional. Is that right?',
    'Session expiry while browsing — do you refresh silently, or do we react? What should the user see, and '
    'what happens to a feed that was scoped to entitlements that have just lapsed? Note wokay resolve '
    'entitlement from the database on every request rather than from a token claim, so a revocation takes '
    'effect on the next request rather than at token expiry.',
    'Who owns the entry-choice screen — us presenting it, you handling the outcome? We have built it as the '
    'Access Gate sheet, screen 03. Its two options are now "through my institution" and "browse free '
    'content", the second being wokay\'s anonymous public feed rather than a personal-account path.',
    'If a signed-in user changes institution, what happens to their session? This affects us, you and '
    't4targaryen — and note that wokay return 403 FORBIDDEN_INSTITUTION_MISMATCH when the path institution '
    'does not match the token, so a stale session surfaces as a specific error we can render.',
    'Loan state is the request we most need answered early. Do you expose a bulk lookup — given forty item '
    'identifiers, which does this user currently have on loan? A page of catalogue results cannot make forty '
    'separate calls, and this state decides whether every card shows Borrow or Read. wokay ship items:batch '
    'for exactly this problem on the metadata side, capped at 100 ids, which is the precedent we would like '
    'to follow. If only a per-item call exists, tell us in Week 1 so we can design the caching around it.',

    'The Borrow endpoint: params, and the complete set of failure responses. We understand from wokay that a '
    'denial is a return value carrying an enumerated reason rather than an exception, with codes '
    'NO_ENTITLEMENT, ENTITLEMENT_EXPIRED, ENTITLEMENT_SUSPENDED, ITEM_ARCHIVED, INSTITUTION_INACTIVE and '
    'CONTENT_NOT_READY. Do those reach us unchanged in the HTTP response, in wokay\'s error envelope? Our '
    'error copy is keyed on the code, so passing it through is what lets us say the library subscription has '
    'expired rather than not available.',

    'The Borrow endpoint URL template — wokay take it from config and it becomes the href in every '
    'acquisition link they emit, which means WE never construct it and never need to know it. Confirming '
    'only that you have given it to them.',

    'On Elite, we understand Borrow is two-outcome: 201 with a loan when a Redis lease is acquired, or 202 '
    'with a queue position when none are free. Confirm the 202 shape, because our button has to flip to a '
    'queue-position state rather than to Read.',

    'GET /api/v1/availability?itemId= — confirm the response shape. We call it on the detail screen only, '
    'Elite only, because wokay\'s feed carries copies.total and never copies.available. This is the only way '
    'we can show a waitlist state before the user taps.',

    'Loan duration and expiry. What happens when a loan lapses while the user is reading, or while they are '
    'offline with a downloaded copy? Note this is Subscription only — Elite writes nothing to the device, so '
    'there is no offline copy to lapse.',

    'TO CONFIRM THE CORRECTED MODEL — and please read this one carefully, because we had it wrong until we '
    'read wokay\'s document. A loan exists for anything ENCRYPTED: Subscription and Elite. A lease and a '
    'queue exist only where copies are FINITE, which is Elite alone. Open access has no entitlement, no '
    'licence and no loan record in ANY session state, including signed in — signing in changes nothing '
    'about it. Subscription is loan-backed but unlimited, so it is downloadable and can never queue. Elite '
    'is the only tier that borrows against a finite pool, and the only one that cannot be downloaded. Is '
    'that right?',

    'For confirmation: Subscribe is GONE, not B2C-only. Individual subscribers are out of scope per wokay\'s '
    'gate decision 7, and the entitlement resolver is keyed by institution. A non-entitled title reached '
    'through pending intent renders as Access Restricted with no action. Please tell us if you expected a '
    'Subscribe path to survive.',
    'Please confirm that Buy, Subscribe and the earlier purchase-adjacent Borrow are all out of scope. We '
    'have removed all three. If a purchase flow is expected to exist, we need to know in Week 1, not Week 3.',
], start=1):
    p = para(space_after=4)
    run(p, f'{i}.  ', bold=True, color=NAVY)
    p.add_run(q)

h3('For t4targaryen')
for i, q in enumerate([
    'What is the contract for launching the reader from our item view? We hand you an item and an intent; '
    'you and flambeau handle the reading-session call, the key and the decryption.',
    'What is the contract for a download, and who owns the download store? Note downloads exist on Open '
    'Access and Subscription only.',
    'Elite offers no Download in our UI, driven by canPersist false on the acquisition link. wokay also '
    'refuse a download intent server-side with DOWNLOAD_NOT_PERMITTED, so the rule is enforced in two places '
    'already. Confirming you hide it at the reader too, so it is three.',
    'THE ONE THAT MATTERS MOST, and it is now a decision rather than a question: who owns screen 08, the '
    'Library / Downloads shelf? Our plan assigns screens 07 and 08 to you as one CAP-7 block, on the '
    'reasoning that everything decision-bearing on screen 08 is yours — download progress, bytes on disk, '
    'the keystore, and guest-download persistence. But wokay\'s section 03 lists team1 as owning "my '
    'library" and pitches their items:batch endpoint at that screen. If it is ours, a ProgressBar component '
    'and a whole screen come back into our Week 1 library and we need to know by Friday. We are adopting '
    'items:batch into our data layer either way, since it is the right shape for any surface holding item '
    'ids without metadata.',
    'Do guest downloads exist — an anonymous user downloading Open Access content — and who stores them? '
    'The design shows a Guest Downloads section on screen 08. Note this is the ONE case where an anonymous '
    'download is possible at all: open access is stored as plaintext with no key and no licence, precisely '
    'so an anonymous reader can open it.',
    'wokay put four properties on every acquisition link for you — encrypted, hasSearchIndex, canPersist and '
    'originalLength. We read canPersist and encrypted; we pass hasSearchIndex and originalLength through '
    'untouched. Confirm you want them from the link we already hold rather than fetching the publication '
    'again.',
], start=1):
    p = para(space_after=4)
    run(p, f'{i}.  ', bold=True, color=NAVY)
    p.add_run(q)

h3('For leadership')
body_para([
    ('Send items 1 and 2 as one letter.', True),
    ' They are the same conversation about the same signed document, and the second is the more urgent of '
    'the two. Both are now backed by wokay\'s published contract rather than by our own preference, which '
    'makes them statements of fact requiring ratification rather than requests for permission.'
])
for i, q in enumerate([
    'PLEASE RATIFY THE CORRECTED ACCESS MATRIX. Design Specification section 3.1 is a signed document and '
    'its access matrix is inverted on two of the three tiers. It states that Open Access requires a Borrow '
    'once the user signs in, and that Elite never borrows and never creates a loan. wokay\'s published '
    'contract states the opposite on both: open access has no entitlement, no licence and no loan record in '
    'any session state, and Elite is the ONLY tier that borrows, consumes a copy and can form a queue. The '
    'error came from the rule "a loan is required for anything the user can take a copy of", which is not '
    'the rule the backend implements — the real rule is that a loan exists for anything encrypted, and a '
    'lease only where copies are finite. We have corrected the specification to v0.3 and are building to it, '
    'because a builder following the v0.2 table would write a resolver that is wrong on most of the '
    'catalogue. We need the correction ratified so that only one matrix is in circulation. This also closes '
    'the open question flagged at v0.2 about signing in making Open Access harder to reach, which was an '
    'artefact of the same error.',

    'Please ratify that the anonymous catalogue is OPEN ACCESS ONLY. This reverses Design Specification '
    'section 4.1 and step 2 of the flow in index.html, both of which state that the catalogue is never '
    'filtered by institution and shows the full catalogue to all users. This is no longer a change we are '
    'proposing — wokay CANNOT SERVE that model. Their feeds are entitlement-scoped and token-protected, and '
    'the unauthenticated endpoint returns open-access titles only. The design package is being corrected to '
    'match, and we need the reversal in writing.',

    'Please confirm that Buy, Subscribe and individual subscribers are all out of the prototype. wokay have '
    'cut individual subscribers at their own gate — their entitlement resolver is keyed by institution and '
    'they judge that changing it adds no capability to the demo. That removes the Subscribe action from our '
    'action vocabulary and removes the personal-account option from our Access Gate screen, which becomes a '
    'browse-free-content entry instead. Screen 13, email sign-in, has no owner under this model.',

    'Please ratify presenting the catalogue as shelf tabs rather than one merged list, which corrects design '
    'screen 01. Note this is narrower than previously described: wokay serve ONE root feed per institution '
    'carrying navigation rows and groups, and the shelf set is configured PER INSTITUTION through their '
    'admin console. So there is nothing to merge and nothing to split — the only question is whether shelves '
    'present as tabs. We are driving the tab bar from data either way, because the backend can re-cut '
    'shelves without a rebuild.',

    'One shared application across the four teams, or one per team? Note wokay state that team1 and '
    't4targaryen have no backend module at all, React Native only, which settles the backend half.',

    'A SCOPE CONFLICT NEEDING AN OWNER, not a ratification. wokay\'s document lists team1 as owning "my '
    'library"; our plan assigns screens 07 and 08 to t4targaryen as one CAP-7 block. It is a whole screen '
    'plus a component, so it changes one person\'s week. We are raising it with both teams and hold that it '
    'is t4targaryen\'s, but if it is unresolved by Friday of Week 1 it needs deciding above us.',
], start=1):
    p = para(space_after=4)
    run(p, f'{i}.  ', bold=True, color=NAVY)
    p.add_run(q)

doc.save(OUT)
print('saved', OUT)
